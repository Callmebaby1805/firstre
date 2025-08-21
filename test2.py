import os
import re
import json
import zipfile
import shutil
import tempfile
from collections import defaultdict

import gradio as gr
import anthropic
from docx import Document
from docx.oxml.shared import qn
from dotenv import load_dotenv
from lxml import etree as ET

# --------------------------- Setup ---------------------------

# Load environment variables
load_dotenv("env (1)")

# Anthropic client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

CATEGORY_TO_JSON = {
    "Sustainability and the Environment": "Sustainabilityandenvironment.json",
    "Strategic Technologies": "StrategicTechnologies.json",  # fixed spelling
    "Trade and Economics": "Trade.json",
    "Politics, Society and Governance": "PSG.json",
    "International Relations, Multipolarity and Multilateralism": "IEMTMLT.json"
}

# -------------------- Quote-aware helpers --------------------

# Treat "…", “…” always as quotes.
# Treat '…' / ‘…’ as quotes ONLY when they wrap a segment (not apostrophes in words like India's).
QUOTE_SPLIT_RE = re.compile(
    r'(".*?"|“.*?”|(?<!\w)\'.*?\'(?!\w)|(?<!\w)‘.*?’(?!\w))',
    re.DOTALL
)

def split_preserve_quotes(text: str):
    """
    Return list of (is_quoted: bool, segment: str).
    Quoted segments are detected and preserved verbatim.
    Apostrophes inside words are NOT treated as quotes.
    """
    parts = []
    tokens = QUOTE_SPLIT_RE.split(text)
    for tok in tokens:
        if not tok:
            continue
        if QUOTE_SPLIT_RE.fullmatch(tok):
            parts.append((True, tok))   # quoted
        else:
            parts.append((False, tok))  # unquoted
    return parts

def apply_outside_quotes(text: str, fn):
    """Apply fn only to unquoted segments and rejoin."""
    out = []
    for is_quoted, seg in split_preserve_quotes(text):
        out.append(seg if is_quoted else fn(seg))
    return "".join(out)

# -------------------- Paragraph run helpers --------------------

def iter_w_t(paragraph):
    """Yield all w:t nodes for a paragraph."""
    return list(paragraph._element.iter(qn('w:t')))

def get_para_text(paragraph):
    """Concatenate paragraph text across all w:t nodes."""
    return "".join(t.text or "" for t in iter_w_t(paragraph))

def set_para_text_preserve(paragraph, new_text):
    """
    Write new_text back into existing w:t nodes by original chunk lengths
    to preserve inline formatting.
    """
    t_nodes = iter_w_t(paragraph)
    if not t_nodes:
        return
    pos = 0
    for i, t in enumerate(t_nodes):
        old = t.text or ""
        if i < len(t_nodes) - 1:
            t.text = new_text[pos:pos+len(old)]
            pos += len(old)
        else:
            t.text = new_text[pos:]

def is_in_table(paragraph) -> bool:
    """Return True if paragraph is inside any table (w:tbl ancestor)."""
    el = paragraph._element
    while el is not None:
        if el.tag == qn('w:tbl'):
            return True
        el = el.getparent()
    return False

def is_heading(paragraph) -> bool:
    """Heuristic: treat built-in/custom headings by style name."""
    try:
        name = paragraph.style.name or ""
    except Exception:
        name = ""
    return name.startswith("Heading")

# ---------------------- DOCX preservation ----------------------

def preserve_extra_parts(original_docx_path, edited_docx_path, final_output_path="expanded_result.docx"):
    """
    Copy only word/document.xml from edited_docx into the original container,
    preserving footnotes, headers, images, styles, etc.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        orig_extract = os.path.join(tmpdir, "orig")
        edit_extract = os.path.join(tmpdir, "edit")
        os.makedirs(orig_extract, exist_ok=True)
        os.makedirs(edit_extract, exist_ok=True)

        with zipfile.ZipFile(original_docx_path, 'r') as z:
            z.extractall(orig_extract)
        with zipfile.ZipFile(edited_docx_path, 'r') as z:
            z.extractall(edit_extract)

        # Replace only the main document part
        shutil.copyfile(
            os.path.join(edit_extract, "word", "document.xml"),
            os.path.join(orig_extract, "word", "document.xml")
        )

        # Rezip to final file
        with zipfile.ZipFile(final_output_path, 'w', zipfile.ZIP_DEFLATED) as outzip:
            for folder, _, files in os.walk(orig_extract):
                for f in files:
                    p = os.path.join(folder, f)
                    arc = os.path.relpath(p, orig_extract)
                    outzip.write(p, arc)
    return final_output_path

# --------------------- Core text utilities ---------------------

def normalize_key(text):
    return re.sub(r'\s+', '', text.lower())

def clean_nested_expansion(text):
    # Term (Term (ABBR)) -> Term (ABBR),  Term (Term) -> Term
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

def handle_special_case_usa(text):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)
    return text

def resolve_synonyms_first_mention_global(text, synonym_groups):
    """
    For each group, the first term that appears wins; replace other synonyms with it (case-insensitive).
    """
    for group in synonym_groups:
        positions = {}
        for term in group:
            m = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
            if m:
                positions[term] = m.start()
        if positions:
            chosen = min(positions, key=positions.get)
            for alt in group:
                if alt.lower() != chosen.lower():
                    pattern = rf'\b{re.escape(alt)}\b'
                    text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)
    return text

# --------------------- Percentage transforms ---------------------

PERCENT_SYMBOL_RE = re.compile(r'\b(\d+(?:\.\d+)?)\s*%')          # e.g., 5% or 5.5%
WORD_PERCENT_RE   = re.compile(r'\bpercent\b', flags=re.IGNORECASE)  # 'percent' -> 'per cent'
WORD_PER_CENT_RE  = re.compile(r'\bper\s*[--]?\s*cent\b', flags=re.IGNORECASE)  # normalise variants

def transform_percentages(segment: str) -> str:
    """
    Enforce percentage rule in unquoted text:
    - 5%  -> 5 per cent
    - percent -> per cent
    - normalise 'per-cent' -> 'per cent'
    """
    segment = PERCENT_SYMBOL_RE.sub(lambda m: f"{m.group(1)} per cent", segment)
    segment = WORD_PERCENT_RE.sub("per cent", segment)
    segment = WORD_PER_CENT_RE.sub("per cent", segment)
    return segment

# --------------------- Topic classifier ---------------------

def classify_topic(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    paragraph = paragraphs[1] if len(paragraphs) > 1 else ""

    categories = list(CATEGORY_TO_JSON.keys())
    system_prompt = (
        "You are an expert in classifying documents. Based only on the given title and paragraph, "
        f"classify the content into one of the following categories:\n{', '.join(categories)}.\n"
        "Respond ONLY with the exact name of the category."
    )
    user_prompt = f"Title: {title}\nParagraph: {paragraph}"

    response = client.messages.create(
        model="claude-3-5-haiku-20241022",
        max_tokens=100,
        temperature=0,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}]
    )

    return response.content[0].text.strip()

# -------------------- Main processing pipeline --------------------

def expand_and_apply_percentages(doc_path, base_json_path="abbreviations.json"):
    # Load base & category maps
    with open(base_json_path, "r", encoding="utf-8") as f:
        base_map = json.load(f)

    category = classify_topic(doc_path)
    category_map = {}
    category_json = CATEGORY_TO_JSON.get(category)
    if category_json and os.path.exists(category_json):
        with open(category_json, "r", encoding="utf-8") as f:
            category_map = json.load(f)

    abbr_map = {**base_map, **category_map}
    counts = defaultdict(int)
    norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

    # Open the original doc (we'll edit in place to preserve formatting runs)
    doc = Document(doc_path)

    # Synonym groups (first-mention wins)
    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China", "People's Republic of China", "PRC"],  # fixed wording
    ]

    # -------- Pre-pass: synonym resolution outside quotes on body paras --------
    body_paras = []
    for p in doc.paragraphs:
        if not is_heading(p) and not is_in_table(p):
            body_paras.append(p)

    joined = "\n".join(get_para_text(p) for p in body_paras)

    def _syn_fn(s):
        return resolve_synonyms_first_mention_global(s, synonym_groups)

    full_text = apply_outside_quotes(joined, _syn_fn)
    updated_chunks = full_text.split("\n")

    # Write synonyms back into the same runs to preserve formatting
    idx = 0
    for p in doc.paragraphs:
        if not is_heading(p) and not is_in_table(p):
            set_para_text_preserve(p, updated_chunks[idx])
            idx += 1

    # -------- Main pass: names, abbreviations, USA, percentages (outside quotes) --------
    name_counts = defaultdict(int)
    name_pattern = re.compile(r'\b(?:Mr|Mrs|Ms|Dr|Prof)?\.?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b')

    def replace_names(segment: str) -> str:
        matches = list(name_pattern.finditer(segment))
        # Replace sequentially using match spans to avoid global collisions
        offset = 0
        for m in matches:
            full = m.group(1).strip()
            name_counts[full] += 1
            if name_counts[full] > 1:
                last = full.split()[-1]
                # Compute current span with offset
                s, e = m.start(), m.end()
                s += offset; e += offset
                segment = segment[:s] + segment[s:e].replace(m.group(0), last, 1) + segment[e:]
                # Adjust offset by change in length
                delta = len(last) - len(m.group(0))
                offset += delta
        return segment

    def process_segment(seg: str) -> str:
        """
        Runs on UNQUOTED text only.
        Applies: name condensation -> clean nested expansions -> first-mention Expansion (ABBR),
        subsequent ABBR -> USA special case -> percentage rule.
        Relies on outer-scope: re, counts, norm_map, replace_names, clean_nested_expansion,
        handle_special_case_usa, transform_percentages.
        """
        # 1) Name condensation (after first full-name mention -> last name)
        seg = replace_names(seg)

        # 2) Clean nested expansions like "Term (Term (ABBR))" -> "Term (ABBR)"
        seg = clean_nested_expansion(seg)

        # 3) Abbreviation expansion
        for _, (abbr, expansion) in norm_map.items():
            abbr_b = rf'\b{re.escape(abbr)}\b'
            exp_b  = rf'\b{re.escape(expansion)}\b'
            patterns = [
                rf'{abbr_b}\s*\(\s*{exp_b}\s*\)',  # ABBR (EXP)
                rf'{exp_b}\s*\(\s*{abbr_b}\s*\)',  # EXP (ABBR)
                rf'{abbr_b}\s*\(\s*[^)]+\s*\)',    # ABBR (something)
                rf'[^()]*\(\s*{abbr_b}\s*\)',      # something (ABBR)
                abbr_b,                             # ABBR
                exp_b                               # EXP
            ]

            if any(re.search(pat, seg, re.IGNORECASE) for pat in patterns):
                counts[abbr] += 1
                if counts[abbr] == 1:
                    # First mention must be "Expansion (ABBR)"
                    correct = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
                    if not re.search(correct, seg, re.IGNORECASE):
                        seg = re.sub('|'.join(patterns), f"{expansion} ({abbr})",
                                     seg, count=1, flags=re.IGNORECASE)
                else:
                    # Subsequent mentions -> ABBR only
                    seg = re.sub('|'.join(patterns), abbr, seg, flags=re.IGNORECASE)

        # 4) USA special-case normalisation
        seg = handle_special_case_usa(seg)

        # 5) Percentage rule for body paragraphs (not tables)
        seg = transform_percentages(seg)

        return seg

    for p in doc.paragraphs:
        if is_in_table(p):
            # Skip tables entirely (keep % in tables/figures)
            continue

        original = get_para_text(p)
        if not original:
            continue

        updated = apply_outside_quotes(original, process_segment)

        if updated != original:
            set_para_text_preserve(p, updated)

    # Save to a temp edited file, then swap only document.xml back into original container
    tmp_edited = "tmp_modified.docx"
    doc.save(tmp_edited)
    final_path = preserve_extra_parts(original_docx_path=doc_path, edited_docx_path=tmp_edited,
                                      final_output_path="expanded_result.docx")
    return final_path

# ------------------------- Gradio UI -------------------------

def gradio_process(doc_file):
    """
    Single-button UI: applies abbreviation expansion + percentage rule.
    No checkboxes / feature toggles.
    """
    out = expand_and_apply_percentages(doc_file.name)
    return "✅ Process complete (abbreviations + percentage rule applied).", out

gr.Interface(
    fn=gradio_process,
    inputs=[gr.File(label="Upload DOCX Document", file_types=[".docx"])],
    outputs=[gr.Text(label="Status"), gr.File(label="Download Processed File")],
    title="ISAS Cleaner: Abbreviations + Percentage Rule",
    description="Applies first-mention abbreviation formatting, synonym normalisation (outside quotes), and converts body-paragraph % to 'per cent' (tables untouched)."
).launch()
