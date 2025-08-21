import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

class DocumentProcessor:
    """A class to process Word documents and apply font, size, and color changes while preserving structure."""
    
    def __init__(self, input_path, output_path):
        """Initialize the processor with input and output file paths."""
        self.input_path = input_path
        self.output_path = output_path
        self.doc = Document(input_path)
    
    def set_font_and_size(self, run, font_name="Calibri", size=None):
        """Set the font and size for a given run, preserving other formatting."""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        if size is not None:
            run.font.size = Pt(size)
        print(f"Set run with text '{run.text[:30]}...' to font: {font_name}, size: {size if size else 'unchanged'}")
    
    def set_paragraph_style_font(self, paragraph, font_name="Calibri"):
        """Set the font for a paragraph's style to ensure consistency."""
        if paragraph.style and paragraph.style.font:
            paragraph.style.font.name = font_name
            paragraph.style.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            paragraph.style.font._element.rPr.rFonts.set(qn('w:ascii'), font_name)
            paragraph.style.font._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        print(f"Set paragraph style font for '{paragraph.text.strip()[:30]}...' to {font_name}")
    
    def set_text_color(self, run):
        """Set the text color of a run to black."""
        run.font.color.rgb = RGBColor(0, 0, 0)
        print(f"Set run with text '{run.text[:30]}...' to black color")
    
    def is_main_heading(self, paragraph, is_first_non_empty):
        """Identify if a paragraph is the main heading (first non-empty paragraph)."""
        has_text = bool(paragraph.text.strip())
        result = is_first_non_empty and has_text
        print(f"Checking paragraph: '{paragraph.text.strip()[:30]}...', is_first_non_empty: {is_first_non_empty}, has_text: {has_text}, is_main_heading: {result}")
        return result
    
    def is_toc_paragraph(self, paragraph, is_after_contents):
        """Check if a paragraph is part of the Table of Contents."""
        text = paragraph.text.strip().lower()
        if text == 'contents':
            return True
        if is_after_contents:
            return bool(text)  # Non-empty paragraphs after "Contents" are TOC
        return False
    
    def is_section_heading(self, text):
        """Check if a paragraph is a section heading (not a TOC entry)."""
        text = text.strip().lower()
        section_headings = {'abbreviations', 'executive summary', 'introduction', 'conclusion'}
        return text in section_headings
    
    def process_paragraphs(self):
        """Process all paragraphs in the document, including Table of Contents."""
        first_non_empty_found = False
        is_after_contents = False
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():  # Skip empty paragraphs
                print(f"Skipping empty paragraph {i}")
                continue
            text = para.text.strip()
            print(f"Processing paragraph {i}: '{text[:30]}...' with {len(para.runs)} runs")
            if text.lower() == 'contents':
                is_after_contents = True
            is_main = self.is_main_heading(para, not first_non_empty_found)
            is_toc = self.is_toc_paragraph(para, is_after_contents)
            for run in para.runs:
                self.set_font_and_size(run, size=14 if is_main else 12)
                self.set_text_color(run)
            if is_toc:
                self.set_paragraph_style_font(para)
                print(f"Processed TOC paragraph: '{text[:30]}...'")
            if is_main:
                first_non_empty_found = True
                print(f"Applied 14 pt to main heading: '{text[:30]}...'")
            if is_after_contents and self.is_section_heading(text):
                is_after_contents = False
    
    def process_tables(self):
        """Process all tables in the document."""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():  # Skip empty paragraphs
                            continue
                        print(f"Processing table paragraph: '{para.text.strip()[:30]}...'")
                        for run in para.runs:
                            self.set_font_and_size(run, size=12)
                            self.set_text_color(run)
    
    def process_footnotes(self):
        """Process all footnotes in the document."""
        if hasattr(self.doc, 'footnotes'):
            for footnote in self.doc.footnotes.footnotes:
                for para in footnote.paragraphs:
                    if not para.text.strip():  # Skip empty paragraphs
                        continue
                    print(f"Processing footnote paragraph: '{para.text.strip()[:30]}...'")
                    for run in para.runs:
                        self.set_font_and_size(run, size=10)
                        self.set_text_color(run)
    
    def process_headers_footers(self):
        """Process headers and footers in the document."""
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                for para in header.paragraphs:
                    if not para.text.strip():  # Skip empty paragraphs
                        continue
                    print(f"Processing header paragraph: '{para.text.strip()[:30]}...'")
                    for run in para.runs:
                        self.set_font_and_size(run, size=12)
                        self.set_text_color(run)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                for para in footer.paragraphs:
                    if not para.text.strip():  # Skip empty paragraphs
                        continue
                    print(f"Processing footer paragraph: '{para.text.strip()[:30]}...'")
                    for run in para.runs:
                        self.set_font_and_size(run, size=12)
                        self.set_text_color(run)
    
    def process(self):
        """Process the entire document and save the output."""
        print(f"Starting processing for {self.input_path}")
        # Process all document components
        self.process_paragraphs()
        self.process_tables()
        self.process_footnotes()
        self.process_headers_footers()
        
        # Save the processed document
        self.doc.save(self.output_path)
        print(f"Document processed and saved to {self.output_path}")

if __name__ == "__main__":
    # Example usage
    input_file = "South10Page.docx"
    output_file = "docpart1.docx"
    processor = DocumentProcessor(input_file, output_file)
    processor.process()