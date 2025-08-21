# app.py
# Gradio wrapper for your style rules with .docx-safe handling.
# - If input is .docx: open it with python-docx and transform ONLY body paragraphs not inside tables.
#   (Headers, footers, footnotes, images, tables remain untouched.)
# - If input is text/markdown/html: process as text and export to .docx.
# - Rules: numbers (0–9 to words), decimals (to words unless part of % or temperature),
#          N% / N percent -> N per cent, currencies (GBP/EUR/JPY/SGD/USD) code->symbol,
#          temperatures: "1.5°C" -> "1.5 Degree Celsius", "98°F" -> "98 Degree Fahrenheit".
# - XML sanitization prevents invalid characters in the .docx.

import os
import io
import time
import re
import unicodedata
import gradio as gr

try:
    from docx import Document
except ImportError:
    Document = None  # we'll error nicely if missing

# ------------------------
# Helpers: masking & scope
# ------------------------

def _merge_spans(spans):
    if not spans:
        return []
    spans = sorted(spans)
    merged = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return [(s, e) for s, e in merged]

def _mask_ranges(text, spans, tag):
    spans = _merge_spans(spans)
    buckets, out, last = [], [], 0
    for s, e in spans:
        out.append(text[last:s])
        out.append(f'§{tag}§{len(buckets)}§')
        buckets.append(text[s:e])
        last = e
    out.append(text[last:])
    return ''.join(out), buckets

def _unmask(text, tag, buckets):
    for i, v in enumerate(buckets):
        text = text.replace(f'§{tag}§{i}§', v)
    return text

def _find_code_fences(text):
    return [(m.start(), m.end()) for m in re.finditer(r'```.*?\n.*?```', text, flags=re.DOTALL)]

def _find_quoted_segments(text):
    pats = [r'"[^"\n]*"', r"'[^'\n]*'", r'“[^”\n]*”', r'‘[^’\n]*’']
    spans = []
    for p in pats:
        spans += [(m.start(), m.end()) for m in re.finditer(p, text)]
    return _merge_spans(spans)

def _is_table_or_figure_line(line):
    if line.count('|') >= 2: return True
    if line.count('\t') >= 2: return True
    if re.search(r'^\s*(Table|Figure|Fig\.|Chart|Exhibit)\b', line, re.IGNORECASE): return True
    if re.match(r'^\s*\|? *-+ *(?:\| *-+ *)+\|?\s*$', line): return True
    return False

def _find_html_blocks(text, tags):
    spans = []
    for tag in tags:
        pattern = re.compile(rf'<{tag}\b[^>]*>.*?</{tag}>', re.IGNORECASE | re.DOTALL)
        spans += [(m.start(), m.end()) for m in pattern.finditer(text)]
    return spans

def _find_html_single_tags(text, tag):
    pat = re.compile(rf'<{tag}\b[^>]*>', re.IGNORECASE | re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_heading_spans(text):
    spans = []
    for m in re.finditer(r'(?m)^(#{1,6})\s.*$', text): spans.append((m.start(), m.end()))
    for m in re.finditer(r'(?ms)^(?P<h>.+?)\n(=+|-{3,})\s*$', text): spans.append((m.start('h'), m.end()))
    return spans

def _find_md_image_spans(text):
    pat = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_footnote_block_spans(text):
    spans = []
    for m in re.finditer(r'(?m)^\[\^[^\]]+\]:[^\n]*\n(?:[ \t].*\n)*', text): spans.append((m.start(), m.end()))
    return spans

def _find_docx_placeholder_spans(text):
    pats = [r'§FOOTNOTE:[^§]+§', r'FN_PLACEHOLDER[^§\n]*']
    spans = []
    for p in pats:
        for m in re.finditer(p, text):
            spans.append((m.start(), m.end()))
    return spans

def _collect_protected_spans(text):
    spans = []
    spans += _find_code_fences(text)
    spans += _find_quoted_segments(text)
    spans += _find_html_blocks(text, ['table','thead','tbody','tfoot','tr','header','footer','figure'])
    spans += _find_html_single_tags(text, 'img')
    spans += _find_md_heading_spans(text)
    spans += _find_md_image_spans(text)
    spans += _find_md_footnote_block_spans(text)
    spans += _find_docx_placeholder_spans(text)
    return _merge_spans(spans)

# ------------------------
# Currency normalization
# ------------------------

def _apply_currency_rules(text: str) -> str:
    CODES = {'GBP':'£','EUR':'€','JPY':'¥','SGD':'S$','USD':'US$'}
    for code, sym in CODES.items():
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)\b', rf'{sym}\1', text, flags=re.IGNORECASE)
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)\s*(million|billion)\b', rf'{sym}\1 \2', text, flags=re.IGNORECASE)
    text = re.sub(r'(?:(?<=€)|(?<=£)|(?<=¥)|(?<=₹)|(?<=S\$)|(?<=US\$))\s+(?=\d)', '', text)
    text = re.sub(r'((?:€|£|¥|₹|S\$|US\$)\d[\d,]*(?:\.\d+)?)\s*(million|billion)\b', r'\1 \2', text, flags=re.IGNORECASE)
    return text

def _mask_currencies(text):
    pat = re.compile(r'(?:US\$|S\$|€|£|¥|₹)\s?\d[\d,]*(?:\.\d+)?(?:\s+(?:million|billion))?', re.IGNORECASE)
    spans = [(m.start(), m.end()) for m in pat.finditer(text)]
    return _mask_ranges(text, spans, 'CUR')

# ------------------------
# Numbers to words (integers & decimals)
# ------------------------

_WORDS_0_19 = ["zero","one","two","three","four","five","six","seven","eight","nine",
               "ten","eleven","twelve","thirteen","fourteen","fifteen","sixteen",
               "seventeen","eighteen","nineteen"]
_TENS = ["","","twenty","thirty","forty","fifty","sixty","seventy","eighty","ninety"]
_SCALES = [(10**9,"billion"), (10**6,"million"), (1000,"thousand"), (100,"hundred")]

def _int_to_words(n: int) -> str:
    if n < 20: return _WORDS_0_19[n]
    if n < 100:
        tens, rem = divmod(n, 10)
        return _TENS[tens] + ("" if rem == 0 else "-" + _WORDS_0_19[rem])
    for value, name in _SCALES:
        if n >= value:
            major, rem = divmod(n, value)
            left = _int_to_words(major) + f" {name}"
            if rem == 0: return left
            return left + " " + _int_to_words(rem)
    return str(n)

def _decimal_to_words_token(tok: str) -> str:
    s = tok.strip()
    neg = s.startswith('-') or s.startswith('−')
    if neg: s = s[1:]
    if s.startswith('.'): int_part, frac_part = 0, s[1:]
    else: int_part, frac_part = s.split('.', 1)
    int_words = _int_to_words(int(int_part)) if int_part else "zero"
    frac_words = " ".join(_WORDS_0_19[int(d)] for d in frac_part if d.isdigit())
    return f"{'minus ' if neg else ''}{int_words} point {frac_words}"

def _wordnum_to_int(word: str) -> int:
    m = {"zero":0,"one":1,"two":2,"three":3,"four":4,"five":5,"six":6,"seven":7,"eight":8,"nine":9,
         "ten":10,"eleven":11,"twelve":12,"thirteen":13,"fourteen":14,"fifteen":15,"sixteen":16,
         "seventeen":17,"eighteen":18,"nineteen":19,
         "twenty":20,"thirty":30,"forty":40,"fifty":50,"sixty":60,"seventy":70,"eighty":80,"ninety":90}
    w = word.lower()
    if '-' in w:
        a, b = w.split('-', 1)
        return m[a] + m[b]
    return m[w]

# ------------------------
# Core: combined transformer
# ------------------------

def apply_all_rules(text: str) -> str:
    """
    Combines: numbers, decimals, percentages, currency, temperature.
    Protections: quotes, code fences, tables, images, headers, footers, footnotes are untouched (for text inputs).
    """
    text = _apply_currency_rules(text)

    prot_spans = _collect_protected_spans(text)
    text, prot_bucket = _mask_ranges(text, prot_spans, 'PROT')

    text, cur_bucket = _mask_currencies(text)

    def _is_scientific_context(line, start, end):
        ctx = line[max(0, start-6): min(len(line), end+6)]
        return bool(re.search(r'[=≈±]', ctx))

    lines = text.splitlines(keepends=False)
    for i, line in enumerate(lines):
        if _is_table_or_figure_line(line):
            lines[i] = line
            continue

        # Percentages
        line = re.sub(r'\b(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\.\d+)\s*%\b', r'\1 per cent', line)
        line = re.sub(r'\b(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\.\d+)\s*percent\b(?=[\s\.,;:!\?\)\]]|$)',
                      r'\1 per cent', line, flags=re.IGNORECASE)
        line = re.sub(
            r'\b('
            r'zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|'
            r'fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|'
            r'eighty|ninety(?:-(?:one|two|three|four|five|six|seven|eight|nine))?'
            r')\s*percent\b(?=[\s\.,;:!\?\)\]]|$)',
            lambda m: f"{_wordnum_to_int(m.group(1))} per cent",
            line, flags=re.IGNORECASE
        )

        # Temperature
        def temp_deg_repl(m):
            start, end = m.span()
            if _is_scientific_context(line, start, end): return m.group(0)
            val, unit = m.group(1), m.group(2).upper()
            return f"{val} Degree {'Celsius' if unit=='C' else 'Fahrenheit'}"
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*°\s*([CF])\b', temp_deg_repl, line)

        def temp_sym_repl(m):
            start, end = m.span()
            if _is_scientific_context(line, start, end): return m.group(0)
            val, unit = m.group(1), m.group(2).upper()
            return f"{val} Degree {'Celsius' if unit=='C' else 'Fahrenheit'}"
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*([CF])\b(?![-\d])', temp_sym_repl, line)

        # Decimals (not part of % or temperature)
        def dec_repl(m):
            tok = m.group(1)
            after = line[m.end():]
            if re.match(r'^\s*(?:per\s*cent|percent|%)\b', after, flags=re.IGNORECASE): return tok
            if re.match(r'^\s*(?:°\s*[CF]|[CF]\b|degree\s+(?:celsius|fahrenheit))', after, flags=re.IGNORECASE): return tok
            return _decimal_to_words_token(tok)
        line = re.sub(r'(?<![\w-])((?:\d+\.\d+)|(?:\.\d+))(?!\.\d)(?![\w-])', dec_repl, line)

        # Integers 0–9 -> words (standalone)
        line = re.sub(r'(?<![\w./:-])([0-9])(?![\w./:-])',
                      lambda m: _WORDS_0_19[int(m.group(1))], line)
        # Capitalize if sentence-start
        line = re.sub(r'^(\s*)(zero|one|two|three|four|five|six|seven|eight|nine)\b',
                      lambda m: m.group(1) + m.group(2).capitalize(), line)

        lines[i] = line

    text = '\n'.join(lines)

    text = _unmask(text, 'CUR', cur_bucket)
    text = _unmask(text, 'PROT', prot_bucket)
    return text

# ------------------------
# XML-safe sanitizer for DOCX
# ------------------------

def _xml_safe_text(s: str) -> str:
    if s is None: return ""
    s = unicodedata.normalize("NFC", s)
    s = s.replace("\x0b", "\n").replace("\x0c", "\n")
    s = re.sub(r"[\x00-\x08\x0E-\x1F]", "", s)
    s = re.sub(r"[\uFFFE\uFFFF]", "", s)
    return s

# ------------------------
# .docx handling (in-place transform preserving formatting)
# ------------------------

def _para_in_table(paragraph) -> bool:
    # Works across python-docx versions: uses built-in namespace map
    return bool(paragraph._element.xpath("ancestor::w:tbl"))


def _transform_paragraph_runs(paragraph, transform_fn):
    """Apply transform_fn to the concatenated text of runs and write back across the same runs,
    preserving run formatting (bold/italic/etc)."""
    runs = paragraph.runs
    if not runs:
        return
    original_texts = [r.text for r in runs]
    concat = "".join(original_texts)
    if not concat.strip():
        return
    new_text = transform_fn(concat)

  
    pos = 0
    for idx, r in enumerate(runs):
        if idx < len(runs) - 1:
            take = len(original_texts[idx])
            piece = new_text[pos:pos + take]
            r.text = _xml_safe_text(piece)
            pos += take
        else:
            r.text = _xml_safe_text(new_text[pos:])

def process_docx_file(src_path: str, dst_path: str):
    doc = Document(src_path)
    # Process ONLY body paragraphs not in tables
    for p in doc.paragraphs:
        if _para_in_table(p):
            continue  # leave tables untouched
        _transform_paragraph_runs(p, apply_all_rules)

    # (Headers/footers/footnotes are left as-is; python-docx retains them on save)
    doc.save(dst_path)

def extract_preview_from_docx(path: str, max_lines: int = 80) -> str:
    try:
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                lines.append(txt)
            if len(lines) >= max_lines:
                break
        return "\n".join(lines)
    except Exception:
        return "(Preview unavailable)"

# ------------------------
# Text I/O for non-docx
# ------------------------

def read_text_from_path(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="ignore")

def save_text_as_docx(text: str, out_path: str):
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    from docx import Document as _Doc
    doc = _Doc()
    lines = text.splitlines() or [""]
    for line in lines:
        doc.add_paragraph(_xml_safe_text(line))
    doc.save(out_path)

# ------------------------
# Gradio glue
# ------------------------

def process_file(file_path: str):
    if not file_path:
        return None, "No file uploaded."
    if Document is None:
        return None, "Missing dependency: python-docx. Install with: pip install python-docx"

    name = os.path.basename(file_path)
    base, ext = os.path.splitext(name)
    ts = int(time.time())
    out_name = f"{base}_styled_{ts}.docx"
    out_dir = "/mnt/data" if os.path.isdir("/mnt/data") else "."
    out_path = os.path.join(out_dir, out_name)

    if ext.lower() == ".docx":
        # Open and transform in place (body paragraphs outside tables)
        process_docx_file(file_path, out_path)
        preview = extract_preview_from_docx(out_path)
    else:
        # Treat as text-like: transform full text and export to .docx
        text = read_text_from_path(file_path)
        out_text = apply_all_rules(text)
        save_text_as_docx(out_text, out_path)
        # Quick preview (first ~80 lines of transformed text)
        preview = "\n".join(out_text.splitlines()[:80])

    return out_path, preview

with gr.Blocks(title="Style Rules Transformer") as demo:
    gr.Markdown(
        "### Style Rules Transformer\n"
        "- Upload **.docx** or text/markdown/HTML.\n"
        "- For **.docx**: transforms only body paragraphs **not inside tables**; headers, footers, footnotes, images, and tables are untouched.\n"
        "- For text: transforms everything and exports to **.docx**.\n"
        "- Includes: numbers, decimals, currency, percentage → *per cent*, temperatures.\n"
        "- Requires: `pip install python-docx`."
    )
    with gr.Row():
        inp = gr.File(label="Upload .docx or text/markdown/HTML", file_count="single", type="filepath")
    with gr.Row():
        btn = gr.Button("Process")
    with gr.Row():
        out_file = gr.File(label="Download processed .docx")
    with gr.Row():
        preview = gr.Textbox(label="Preview (first ~80 lines)", lines=20)

    btn.click(process_file, inputs=[inp], outputs=[out_file, preview])

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)
