import os
import json
import re
import gradio as gr
import anthropic
import zipfile
import shutil
import tempfile
from collections import defaultdict
from docx import Document
from dotenv import load_dotenv
from lxml import etree as ET

# Load environment variables
load_dotenv("env (1)")
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

CATEGORY_TO_JSON = {
    "Sustainability and the Environment": "Sustainabilityandenvironment.json",
    "Strategic Technologies": "StatergicTechnologies.json",
    "Trade and Economics": "Trade.json",
    "Politics, Society and Governance": "PSG.json",
    "International Relations, Multipolarity and Multilateralism": "IEMTMLT.json"
}

def normalize_key(text):
    return re.sub(r'\s+', '', text.lower())

def clean_nested_expansion(text):
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

def handle_special_case_usa(text):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)
    return text

def resolve_synonyms_first_mention_global(text, synonym_groups):
    for group in synonym_groups:
        positions = {}
        for term in group:
            match = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
            if match:
                positions[term] = match.start()
        if positions:
            chosen = min(positions, key=positions.get)
            for alt in group:
                if alt.lower() != chosen.lower():
                    pattern = rf'\b{re.escape(alt)}\b'
                    text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)
    return text

# --- Paragraph text helpers ---

def get_paragraph_text(para):
    return "".join([t.text or "" for t in para._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")])

def set_paragraph_text_preserve(para, new_text):
    t_nodes = list(para._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
    if not t_nodes:
        return
    orig_chunks = [t.text or "" for t in t_nodes]
    pos = 0
    for i, t in enumerate(t_nodes):
        if i < len(orig_chunks) - 1:
            span = len(orig_chunks[i])
            t.text = new_text[pos:pos+span]
            pos += span
        else:
            t.text = new_text[pos:]

# --- ZIP preservation for footnotes and extras ---

def preserve_extra_parts(original_docx_path, edited_docx_path, final_output_path="expanded_result.docx"):
    with tempfile.TemporaryDirectory() as tmpdir:
        orig_extract = os.path.join(tmpdir, "orig")
        edit_extract = os.path.join(tmpdir, "edit")
        os.makedirs(orig_extract)
        os.makedirs(edit_extract)

        with zipfile.ZipFile(original_docx_path, 'r') as zip_ref:
            zip_ref.extractall(orig_extract)
        with zipfile.ZipFile(edited_docx_path, 'r') as zip_ref:
            zip_ref.extractall(edit_extract)

        shutil.copyfile(
            os.path.join(edit_extract, "word", "document.xml"),
            os.path.join(orig_extract, "word", "document.xml")
        )

        with zipfile.ZipFile(final_output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
            for foldername, _, filenames in os.walk(orig_extract):
                for filename in filenames:
                    file_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path, orig_extract)
                    docx.write(file_path, arcname)
    return final_output_path

# --- Count footnote references for debug ---

def count_footnote_refs_in_document_xml(docx_path):
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_bytes = z.read('word/document.xml')
    root = ET.fromstring(xml_bytes)
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return len(root.findall('.//w:footnoteReference', namespaces=NS))

# --- Topic classifier ---

def classify_topic(doc_path):
    doc = Document(doc_path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    title = paras[0] if paras else ""
    paragraph = paras[1] if len(paras) > 1 else ""

    categories = list(CATEGORY_TO_JSON.keys())
    system_prompt = (
        "You are an expert in classifying documents. Based only on the given title and paragraph, "
        f"classify the content into one of the following categories:\n{', '.join(categories)}.\n"
        "Respond ONLY with the exact name of the category."
    )
    user_prompt = f"Title: {title}\nParagraph: {paragraph}"

    response = client.messages.create(
        model="claude-3-5-haiku-20241022",
        max_tokens=100,
        temperature=0,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}]
    )
    return response.content[0].text.strip()

# --- Main abbreviation logic ---

def expand_abbreviations(doc_path, base_json_path="abbreviations.json"):
    with open(base_json_path, "r", encoding="utf-8") as f:
        base_map = json.load(f)

    category = classify_topic(doc_path)
    category_map = {}
    category_json = CATEGORY_TO_JSON.get(category)
    if category_json and os.path.exists(category_json):
        with open(category_json, "r", encoding="utf-8") as f:
            category_map = json.load(f)

    abbr_map = {**base_map, **category_map}
    counts = defaultdict(int)
    norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

    input_doc = Document(doc_path)

    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China", "People Republic of China", "PRC"]
    ]

    editable_paras = []
    original_texts = []
    for p in input_doc.paragraphs:
        txt = get_paragraph_text(p)
        if txt is None:
            txt = ""
        editable_paras.append(p)
        original_texts.append(txt)

    if original_texts:
        joined = "\n".join(original_texts)
        syn_joined = resolve_synonyms_first_mention_global(joined, synonym_groups)
        syn_splits = syn_joined.split("\n")
        if len(syn_splits) == len(editable_paras):
            for p, new_txt in zip(editable_paras, syn_splits):
                set_paragraph_text_preserve(p, new_txt)

    name_counts = defaultdict(int)
    name_pattern = re.compile(r'\b(?:Mr|Mrs|Ms|Dr|Prof)?\.?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b')

    def replace_names(text):
        matches = list(name_pattern.finditer(text))
        for match in matches:
            full_name = match.group(1).strip()
            name_counts[full_name] += 1
            last_name = full_name.split()[-1]
            if name_counts[full_name] > 1:
                text = text.replace(match.group(0), last_name, 1)
        return text

    for p in input_doc.paragraphs:
        line = get_paragraph_text(p)
        if not line:
            continue

        seg = replace_names(line)
        seg = clean_nested_expansion(seg)

        for _, (abbr, expansion) in norm_map.items():
            abbr_pattern = rf'\b{re.escape(abbr)}\b'
            exp_pattern = rf'\b{re.escape(expansion)}\b'
            patterns = [
                rf'{abbr}\s*\(\s*{expansion}\s*\)',
                rf'{expansion}\s*\(\s*{abbr}\s*\)',
                rf'{abbr}\s*\(\s*\w[\w\s]*\)',
                rf'\w[\w\s]*\(\s*{abbr}\s*\)',
                abbr_pattern,
                exp_pattern
            ]
            if any(re.search(pat, seg, re.IGNORECASE) for pat in patterns):
                counts[abbr] += 1
                if counts[abbr] == 1:
                    correct_format = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
                    if not re.search(correct_format, seg, re.IGNORECASE):
                        seg = re.sub('|'.join(patterns), f"{expansion} ({abbr})", seg, count=1, flags=re.IGNORECASE)
                else:
                    seg = re.sub('|'.join(patterns), abbr, seg, flags=re.IGNORECASE)

        seg = handle_special_case_usa(seg)
        set_paragraph_text_preserve(p, seg)

    try:
        before_refs = count_footnote_refs_in_document_xml(doc_path)
    except Exception:
        before_refs = None

    tmp_edited_path = "tmp_modified.docx"
    input_doc.save(tmp_edited_path)

    final_path = preserve_extra_parts(original_docx_path=doc_path, edited_docx_path=tmp_edited_path)

    try:
        after_refs = count_footnote_refs_in_document_xml(final_path)
        if before_refs is not None and after_refs is not None and after_refs < before_refs:
            print(f" Footnote references decreased: before={before_refs}, after={after_refs}")
    except Exception:
        pass

    return final_path

# --- Gradio UI ---

def gradio_abbreviation_expand(doc_file):
    output_path = expand_abbreviations(doc_file.name)
    return " Document processed with topic-specific and global abbreviations.", output_path

gr.Interface(
    fn=gradio_abbreviation_expand,
    inputs=[gr.File(label="Upload DOCX Document", file_types=[".docx"])],
    outputs=[gr.Text(label="Status"), gr.File(label="Download Processed File")],
    title="Abbreviation Expander with Topic-Specific Enhancement",
    description="Expands abbreviations using a global list and a category-specific list (auto-detected using Claude). Preserves inline footnote references and all original parts (footnotes, headers, images)."
).launch()
