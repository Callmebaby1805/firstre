# app.py
# Style Rules Transformer + Mandatory Anthropics (Claude) Post-Edit with MAJOR_PROMPT
# - Deterministic passes: numbers, decimals, percentages, ranges, currency, temperatures, UK spellings (Breame),
#   plus masking to protect quotes/tables/headers/footers/figures/footnotes/*...*/**...**/<<par>>.
# - Claude pass is ALWAYS applied after deterministic edits using MAJOR_PROMPT from updated_prompt.py.
# - DOCX path preserves paragraph formatting by transforming only body paragraphs, grouping text runs.

import os
import time
import re
import unicodedata
import gradio as gr

# ---------- Optional: load Anthropic + dotenv ----------
try:
    from anthropic import Anthropic
except Exception:
    Anthropic = None

try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

# Breame (UK spellings)
try:
    from breame.spelling import get_british_spelling
except Exception:
    get_british_spelling = None

# DOCX
try:
    from docx import Document
except ImportError:
    Document = None

# Your big prompt
from updated_prompt import MAJOR_PROMPT

# ======================================================================================
# ENV LOADER
# ======================================================================================

def _load_env():
    if load_dotenv:
        try:
            load_dotenv()  # .env
        except Exception:
            pass
        # Also try a file literally named "env (1)" if present
        try:
            if os.path.exists("env (1)"):
                load_dotenv("env (1)", override=True)
        except Exception:
            pass

_ANTHROPIC_CLIENT = None

def _get_anthropic_client():
    global _ANTHROPIC_CLIENT
    if _ANTHROPIC_CLIENT is not None:
        return _ANTHROPIC_CLIENT
    if Anthropic is None:
        return None
    _load_env()
    api_key = (
        os.getenv("ANTHROPIC_API_KEY")
        or os.getenv("ANTHROPIC_KEY")
        or os.getenv("ANTHROPIC_APIKEY")
    )
    if not api_key:
        return None
    try:
        _ANTHROPIC_CLIENT = Anthropic(api_key=api_key)
        return _ANTHROPIC_CLIENT
    except Exception:
        return None

# ======================================================================================
# SHARED HELPERS (masking, spans, currencies, numbers, UK spelling)
# ======================================================================================

def _merge_spans(spans):
    if not spans:
        return []
    spans = sorted(spans)
    merged = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return [(s, e) for s, e in merged]

def _idx_to_letters(i: int) -> str:
    if i < 0:
        i = 0
    s = ""
    i += 1
    while i > 0:
        i, r = divmod(i - 1, 26)
        s = chr(ord('a') + r) + s
    return s

def _make_token(tag: str, i: int) -> str:
    return f"§{tag}§{_idx_to_letters(i)}§"

def _mask_ranges(text, spans, tag):
    spans = _merge_spans(spans)
    buckets, out, last = [], [], 0
    for _, (s, e) in enumerate(spans):
        out.append(text[last:s])
        out.append(_make_token(tag, len(buckets)))
        buckets.append(text[s:e])
        last = e
    out.append(text[last:])
    return ''.join(out), buckets

def _unmask(text, tag, buckets):
    for i, v in enumerate(buckets):
        text = text.replace(_make_token(tag, i), v)
    return text

def _find_code_fences(text):
    return [(m.start(), m.end()) for m in re.finditer(r'```.*?\n.*?```', text, flags=re.DOTALL)]

def _find_quoted_segments(text):
    spans = []
    for m in re.finditer(r"(?<!\w)'([^'\n]*?)'(?!\w)", text):  # straight single quotes
        spans.append((m.start(), m.end()))
    for m in re.finditer(r'"[^"\n]*"', text):                  # straight double quotes
        spans.append((m.start(), m.end()))
    for m in re.finditer(r'‘[^’\n]*’', text):                  # curly single quotes
        spans.append((m.start(), m.end()))
    for m in re.finditer(r'“[^”\n]*”', text):                  # curly double quotes
        spans.append((m.start(), m.end()))
    return _merge_spans(spans)

def _is_table_or_figure_line(line):
    if line.count('|') >= 2:
        return True
    if line.count('\t') >= 2:
        return True
    if re.search(r'^\s*(Table|Figure|Fig\.|Chart|Exhibit)\b', line, re.IGNORECASE):
        return True
    if re.match(r'^\s*\|? *-+ *(?:\| *-+ *)+\|?\s*$', line):
        return True
    return False

def _find_html_blocks(text, tags):
    spans = []
    for tag in tags:
        pattern = re.compile(rf'<{tag}\b[^>]*>.*?</{tag}>', re.IGNORECASE | re.DOTALL)
        spans += [(m.start(), m.end()) for m in pattern.finditer(text)]
    return spans

def _find_html_single_tags(text, tag):
    pat = re.compile(rf'<{tag}\b[^>]*>', re.IGNORECASE | re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_heading_spans(text):
    spans = []
    for m in re.finditer(r'(?m)^(#{1,6})\s.*$', text):
        spans.append((m.start(), m.end()))
    for m in re.finditer(r'(?ms)^(?P<h>.+?)\n(=+|-{3,})\s*$', text):
        spans.append((m.start('h'), m.end()))
    return spans

def _find_md_image_spans(text):
    pat = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_footnote_block_spans(text):
    spans = []
    for m in re.finditer(r'(?m)^\[\^[^\]]+\]:[^\n]*\n(?:[ \t].*\n)*', text):
        spans.append((m.start(), m.end()))
    return spans

def _find_docx_placeholder_spans(text):
    pats = [r'§FOOTNOTE:[^§]+§', r'FN_PLACEHOLDER[^§\n]*']
    spans = []
    for p in pats:
        for m in re.finditer(p, text):
            spans.append((m.start(), m.end()))
    return spans

def _find_asterisk_marked_spans(text):
    spans = []
    # **...** first
    for m in re.finditer(r'\*\*.*?\*\*', text, flags=re.DOTALL):
        spans.append((m.start(), m.end()))
    # *...* non-greedy
    for m in re.finditer(r'\*(?!\*)(.*?)\*', text, flags=re.DOTALL):
        spans.append((m.start(), m.end()))
    return _merge_spans(spans)

def _find_par_markers(text):
    spans = []
    for m in re.finditer(re.escape("<<par>>"), text):
        spans.append((m.start(), m.end()))
    return spans

def _collect_protected_spans(text):
    spans = []
    spans += _find_code_fences(text)
    spans += _find_quoted_segments(text)
    spans += _find_html_blocks(text, ['table', 'thead', 'tbody', 'tfoot', 'tr', 'header', 'footer', 'figure'])
    spans += _find_html_single_tags(text, 'img')
    spans += _find_md_heading_spans(text)
    spans += _find_md_image_spans(text)
    spans += _find_md_footnote_block_spans(text)
    spans += _find_docx_placeholder_spans(text)
    return _merge_spans(spans)

# Currencies
def _apply_currency_rules(text: str) -> str:
    CODES = {'GBP':'£','EUR':'€','JPY':'¥','SGD':'S$','USD':'US$'}
    for code, sym in CODES.items():
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)\b', rf'{sym}\1', text, flags=re.IGNORECASE)
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)\s*(million|billion)\b',
                      rf'{sym}\1 \2', text, flags=re.IGNORECASE)
    text = re.sub(r'(?:(?<=€)|(?<=£)|(?<=¥)|(?<=₹)|(?<=S\$)|(?<=US\$))\s+(?=\d)', '', text)
    text = re.sub(r'((?:€|£|¥|₹|S\$|US\$)\d[\d,]*(?:\.\d+)?)\s*(million|billion)\b', r'\1 \2', text, flags=re.IGNORECASE)
    return text

def _mask_currencies(text):
    pat = re.compile(r'(?:US\$|S\$|€|£|¥|₹)\s?\d[\d,]*(?:\.\d+)?(?:\s+(?:million|billion))?', re.IGNORECASE)
    spans = [(m.start(), m.end()) for m in pat.finditer(text)]
    return _mask_ranges(text, spans, 'CUR')

# Number words
_WORDS_0_19 = ["zero","one","two","three","four","five","six","seven","eight","nine",
               "ten","eleven","twelve","thirteen","fourteen","fifteen","sixteen",
               "seventeen","eighteen","nineteen"]
_TENS = ["","","twenty","thirty","forty","fifty","sixty","seventy","eighty","ninety"]
_SCALES = [(10**9,"billion"), (10**6,"million"), (1000,"thousand"), (100,"hundred")]

def _int_to_words(n: int) -> str:
    if n < 20:
        return _WORDS_0_19[n]
    if n < 100:
        tens, rem = divmod(n, 10)
        return _TENS[tens] + ("" if rem == 0 else "-" + _WORDS_0_19[rem])
    for value, name in _SCALES:
        if n >= value:
            major, rem = divmod(n, value)
            left = _int_to_words(major) + f" {name}"
            if rem == 0:
                return left
            return left + " " + _int_to_words(rem)
    return str(n)

def _decimal_to_words_token(tok: str) -> str:
    s = tok.strip()
    neg = s.startswith('-') or s.startswith('−')
    if neg:
        s = s[1:]
    if s.startswith('.'):
        int_part, frac_part = 0, s[1:]
    else:
        int_part, frac_part = s.split('.', 1)
    int_words = _int_to_words(int(int_part)) if int_part else "zero"
    frac_words = " ".join(_WORDS_0_19[int(d)] for d in frac_part if d.isdigit())
    return f"{'minus ' if neg else ''}{int_words} point {frac_words}"

def _wordnum_to_int(word: str) -> int:
    m = {"zero":0,"one":1,"two":2,"three":3,"four":4,"five":5,"six":6,"seven":7,"eight":8,"nine":9,
         "ten":10,"eleven":11,"twelve":12,"thirteen":13,"fourteen":14,"fifteen":15,"sixteen":16,
         "seventeen":17,"eighteen":18,"nineteen":19,
         "twenty":20,"thirty":30,"forty":40,"fifty":50,"sixty":60,"seventy":70,"eighty":80,"ninety":90}
    w = word.lower()
    if '-' in w:
        a, b = w.split('-', 1)
        return m[a] + m[b]
    return m[w]

# Sentence-start integer spelling
_TOKEN_RE = r'(?:§(?:PROT|CUR|AST|PAR)§[a-z]+§)'
_SENT_START_INT = re.compile(
    rf'(^|(?<=[.!?])(?:\s|{_TOKEN_RE})+)'     # start of line OR after .!? then spaces/tokens
    rf'(\d{{1,3}}(?:,\d{{3}})*)\b(?!\.\d)'    # integer (no decimals)
)

def _spell_sentence_start_integers(line: str) -> str:
    def _repl(m):
        prefix = m.group(1)
        num_str = m.group(2)
        n = int(num_str.replace(',', ''))
        words = _int_to_words(n)
        return prefix + words[:1].upper() + words[1:]
    return _SENT_START_INT.sub(_repl, line)

# UK spelling (Breame)
_STAY_IZE = {
    "capsize","capsizes","capsized","capsizing",
    "seize","seizes","seized","seizing",
    "prize","prizes","prized","prizing",
    "size","sizes","sized","sizing",
}
_IZE_SUFFIX_RE = re.compile(r'(?:ize|izes|ized|izing|ization|izations)\b')
_WORD_LOWER_RE = re.compile(r'\b([a-z]+)\b')

def _detect_preserve_ize_pref(raw_text: str) -> bool:
    if not get_british_spelling:
        return False
    ize = len(re.findall(r'\b[a-z]+(?:ize|izes|ized|izing|ization|izations)\b', raw_text))
    ise = len(re.findall(r'\b[a-z]+(?:ise|ises|ised|ising|isation|isations)\b', raw_text))
    return ize >= 5 and ize >= max(8, int(1.8 * max(1, ise)))

def _apply_breame_uk_to_line(line: str, preserve_ize: bool) -> str:
    if get_british_spelling is None:
        return line
    def repl(m):
        w = m.group(1)
        if w in _STAY_IZE:
            return w
        uk = get_british_spelling(w)
        if not uk or uk == w:
            return w
        if preserve_ize and _IZE_SUFFIX_RE.search(w):
            return w
        return uk
    return _WORD_LOWER_RE.sub(repl, line)

# ======================================================================================
# DETERMINISTIC RULES PIPELINE (text)
# ======================================================================================

def apply_all_rules(text: str) -> str:
    text = _apply_currency_rules(text)

    preserve_ize = _detect_preserve_ize_pref(text)

    # Mask protected content
    prot_spans = _collect_protected_spans(text)
    text, prot_bucket = _mask_ranges(text, prot_spans, 'PROT')

    # Mask asterisk-marked spans (*...*, **...**)
    ast_spans = _find_asterisk_marked_spans(text)
    text, ast_bucket = _mask_ranges(text, ast_spans, 'AST')

    # Mask <<par>>
    par_spans = _find_par_markers(text)
    text, par_bucket = _mask_ranges(text, par_spans, 'PAR')

    # Mask currencies (so decimals there are not converted to words)
    text, cur_bucket = _mask_currencies(text)

    def _is_scientific_context(line, start, end):
        ctx = line[max(0, start-6): min(len(line), end+6)]
        return bool(re.search(r'[=≈±]', ctx))

    lines = text.splitlines(keepends=False)
    for i, line in enumerate(lines):
        if _is_table_or_figure_line(line):
            lines[i] = line
            continue

        # UK spelling
        line = _apply_breame_uk_to_line(line, preserve_ize)

        # ---- Ranges to "to" (do before % rules) ----
        def _spell_small(n: int) -> str:
            return _WORDS_0_19[n] if 0 <= n <= 9 else str(n)
        # 3-4% → three to four per cent
        line = re.sub(
            r'(?<![\d.,])(\d{1,3})\s*[-–]\s*(\d{1,3})[\s\u00A0\u202F\u2009\u2007\u2008\u2002-\u200A]*(?:%|\uFF05)(?!\w)',
            lambda m: f"{_spell_small(int(m.group(1)))} to {_spell_small(int(m.group(2)))} per cent",
            line
        )
        # 3-4 → three to four
        line = re.sub(
            r'(?<![\d.,])(\d{1,3})\s*[-–]\s*(\d{1,3})(?!\s*(?:%|\uFF05))',
            lambda m: f"{_spell_small(int(m.group(1)))} to {_spell_small(int(m.group(2)))}",
            line
        )

        # ---- Percentages ----
        # 30%YoY → 30 per cent YoY
        line = re.sub(
            r'(?<![\d.,])(\d{1,3}(?:,\d{3})*)[\s\u00A0\u202F\u2009\u2007\u2008\u2002-\u200A]*(?:%|\uFF05)([A-Za-z])',
            r'\1 per cent \2',
            line
        )
        # N% → N per cent  (but not decimals like 3.25%)
        line = re.sub(
            r'(?<![\d.,])(\d{1,3}(?:,\d{3})*)[\s\u00A0\u202F\u2009\u2007\u2008\u2002-\u200A]*(?:%|\uFF05)(?!\w)',
            r'\1 per cent',
            line
        )
        # N percent → N per cent (decimals allowed)
        line = re.sub(
            r'(?<![\w§])(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\.\d+)[\s\u00A0\u202F\u2009\u2007\u2008\u2002-\u200A]*percent\b(?=[\s\.,;:!\?\)\]»›”’]|$)',
            r'\1 per cent',
            line, flags=re.IGNORECASE
        )
        # word-number percent
        line = re.sub(
            r'\b('
            r'zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|'
            r'fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|'
            r'eighty|ninety(?:-(?:one|two|three|four|five|six|seven|eight|nine))?'
            r')\s*percent\b(?=[\s\.,;:!\?\)\]]|$)',
            lambda m: f"{_wordnum_to_int(m.group(1))} per cent",
            line, flags=re.IGNORECASE
        )

        # ---- Temperature ----
        def temp_deg_repl(m):
            start, end = m.span()
            if _is_scientific_context(line, start, end):
                return m.group(0)
            val, unit = m.group(1), m.group(2).upper()
            return f"{val} Degree {'Celsius' if unit=='C' else 'Fahrenheit'}"
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*°\s*([CF])\b', temp_deg_repl, line)

        def temp_sym_repl(m):
            start, end = m.span()
            if _is_scientific_context(line, start, end):
                return m.group(0)
            val, unit = m.group(1), m.group(2).upper()
            return f"{val} Degree {'Celsius' if unit=='C' else 'Fahrenheit'}"
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*([CF])\b(?![-\d])', temp_sym_repl, line)

        # ---- Decimals (non-%/non-temp) → words ----
        def dec_repl(m):
            tok = m.group(1)
            after = line[m.end():]
            if re.match(
                r'^\s*(?:per\s*cent\b|percent\b|[\u00A0\u202F\u2009\u2007\u2008\u2002-\u200A\s]*(?:%|\uFF05)(?=$|[^0-9A-Za-z]))',
                after, flags=re.IGNORECASE
            ):
                return tok
            if re.match(r'^\s*(?:°\s*[CF]|[CF]\b|degree\s+(?:celsius|fahrenheit))', after, flags=re.IGNORECASE):
                return tok
            return _decimal_to_words_token(tok)
        line = re.sub(
            r'(?<![\w./:\-§])((?:\d+\.\d+)|(?:\.\d+))(?!\.\d)(?![\w./:\-§])',
            dec_repl, line
        )

        # ---- Small integers (0–9) → words ----
        line = re.sub(
            r'(?<![\w./:\-§])([0-9])(?![\w./:\-§])',
            lambda m: _WORDS_0_19[int(m.group(1))],
            line
        )

        # ---- Sentence-start integers → spelled ----
        line = _spell_sentence_start_integers(line)

        lines[i] = line

    text = '\n'.join(lines)

    # Unmask in reverse order of masking
    text = _unmask(text, 'CUR', cur_bucket)
    text = _unmask(text, 'PAR', par_bucket)
    text = _unmask(text, 'AST', ast_bucket)
    text = _unmask(text, 'PROT', prot_bucket)
    return text

# ======================================================================================
# CLAUDE (Anthropic) POST-EDIT
# ======================================================================================

def _extract_corrected_text(raw: str) -> str:
    """
    If the model returns a 'Corrected:' header, return only the content after it.
    Otherwise, return raw.
    """
    if raw is None:
        return ""
    m = re.search(r'^\s*Corrected:\s*\n?', raw, flags=re.IGNORECASE | re.MULTILINE)
    if not m:
        return raw.strip()
    return raw[m.end():].strip()

def _anthropic_complete(prompt_text: str, model: str, max_tokens: int):
    client = _get_anthropic_client()
    if client is None:
        return None, "Anthropic client not available (missing SDK or API key)."
    try:
        msg = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            temperature=0,
            messages=[{"role": "user", "content": prompt_text}]
        )
        out = "".join(block.text for block in msg.content if getattr(block, "type", "text") == "text")
        return out, None
    except Exception as e:
        return None, f"Anthropic error: {e}"
def apply_major_prompt_with_masks(
    text: str, model: str, max_tokens: int
) -> tuple[str, str | None]:

    """
    Mask protected content (quotes, code fences, tables/figures/headers/images, *...*, **...**, <<par>>),
    send to Claude using MAJOR_PROMPT, then unmask.
    Returns (result_text, note)
    """
    # Mask
    prot_spans = _collect_protected_spans(text)
    masked, prot_bucket = _mask_ranges(text, prot_spans, 'PROT')

    ast_spans = _find_asterisk_marked_spans(masked)
    masked, ast_bucket = _mask_ranges(masked, ast_spans, 'AST')

    par_spans = _find_par_markers(masked)
    masked, par_bucket = _mask_ranges(masked, par_spans, 'PAR')

    # Compose prompt
    prompt_text = MAJOR_PROMPT.replace("{text}", masked)

    # Call Claude
    out, err = _anthropic_complete(prompt_text, model=model, max_tokens=max_tokens)
    if err or out is None:
        # On any error, return original text and the note
        return text, err or "Unknown Anthropic error."

    # Extract corrected content (if a 'Corrected:' header is present)
    out = _extract_corrected_text(out)

    # Unmask
    out = _unmask(out, 'PAR', par_bucket)
    out = _unmask(out, 'AST', ast_bucket)
    out = _unmask(out, 'PROT', prot_bucket)
    return out, None

# ======================================================================================
# DOCX PROCESSING (group contiguous runs; Claude per-paragraph ALWAYS)
# ======================================================================================

def _xml_safe_text(s: str) -> str:
    if s is None:
        return ""
    s = unicodedata.normalize("NFC", s)
    s = s.replace("\x0b", "\n").replace("\x0c", "\n")
    s = re.sub(r"[\x00-\x08\x0E-\x1F]", "", s)
    s = re.sub(r"[\uFFFE\uFFFF]", "", s)
    return s

def _para_in_table(paragraph) -> bool:
    return bool(paragraph._element.xpath("ancestor::*[local-name()='tbl']"))

def _is_heading(paragraph) -> bool:
    try:
        name = (paragraph.style.name or "").lower()
    except Exception:
        name = ""
    return name.startswith("heading") or name in {"title", "subtitle"}

def _is_quote_style(paragraph) -> bool:
    try:
        name = (paragraph.style.name or "").lower()
    except Exception:
        name = ""
    return "quote" in name

def _run_has_nontext_children(run) -> bool:
    r = run._r
    return bool(r.xpath(".//*[local-name()='footnoteReference' or local-name()='endnoteReference' or local-name()='fldChar' or local-name()='instrText' or local-name()='drawing']"))

def _transform_paragraph_runs_grouped(paragraph, transform_fn):
    runs = paragraph.runs
    n = len(runs)
    i = 0
    while i < n:
        if _run_has_nontext_children(runs[i]):
            i += 1
            continue
        group_idx = []
        group_txt = []
        while i < n and not _run_has_nontext_children(runs[i]):
            t = runs[i].text or ""
            group_idx.append(i)
            group_txt.append(t)
            i += 1
        combined = "".join(group_txt)
        if combined.strip() == "":
            continue
        new_combined = transform_fn(combined)
        if new_combined == combined:
            continue
        pos = 0
        for k, ridx in enumerate(group_idx):
            orig_len = len(group_txt[k])
            if k < len(group_idx) - 1:
                piece = new_combined[pos:pos + orig_len]
                runs[ridx].text = _xml_safe_text(piece)
                pos += orig_len
            else:
                runs[ridx].text = _xml_safe_text(new_combined[pos:])

def process_docx_file(src_path: str, dst_path: str, model: str, max_tokens: int, preview_lines: int = 80):
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    doc = Document(src_path)
    for p in doc.paragraphs:
        # leave tables/headings/quote-styled paragraphs unchanged
        if _para_in_table(p) or _is_heading(p) or _is_quote_style(p):
            continue

        # 1) Deterministic rules
        _transform_paragraph_runs_grouped(p, apply_all_rules)

        # 2) Claude per paragraph (mandatory)
        para_text = p.text
        if para_text.strip():
            edited, err = apply_major_prompt_with_masks(para_text, model=model, max_tokens=max_tokens)
            if not err and edited and edited != para_text:
                for r in p.runs:
                    r.text = ""
                p.add_run(_xml_safe_text(edited))

    doc.save(dst_path)

    # Build preview
    try:
        lines = []
        for p in Document(dst_path).paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
            if len(lines) >= preview_lines:
                break
        return "\n".join(lines)
    except Exception:
        return "(Preview unavailable)"

# ======================================================================================
# TEXT I/O (non-docx)
# ======================================================================================

def read_text_from_path(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="ignore")

def save_text_as_docx(text: str, out_path: str):
    if Document is None:
        raise ImportError("python-docx is not installed. Install with: pip install python-docx")
    from docx import Document as _Doc
    doc = _Doc()
    lines = text.splitlines() or [""]
    for line in lines:
        doc.add_paragraph(_xml_safe_text(line))
    doc.save(out_path)

# ======================================================================================
# GRADIO GLUE
# ======================================================================================

def process_file(file_path: str, model: str, max_tokens: int):
    if not file_path:
        return None, "No file uploaded."
    if Document is None:
        return None, "Missing dependency: python-docx. Install with: pip install python-docx"

    name = os.path.basename(file_path)
    base, ext = os.path.splitext(name)
    ts = int(time.time())
    out_name = f"{base}_styled_{ts}.docx"
    out_dir = "/mnt/data" if os.path.isdir("/mnt/data") else "."
    out_path = os.path.join(out_dir, out_name)

    note_lines = []

    if ext.lower() == ".docx":
        preview = process_docx_file(
            file_path, out_path,
            model=model,
            max_tokens=max_tokens
        )
    else:
        text = read_text_from_path(file_path)
        out_text = apply_all_rules(text)

        edited, err = apply_major_prompt_with_masks(out_text, model=model, max_tokens=max_tokens)
        if err:
            note_lines.append(f"[Note] Claude step failed: {err}")
        else:
            out_text = edited

        save_text_as_docx(out_text, out_path)
        preview = "\n".join(out_text.splitlines()[:80])

    if get_british_spelling is None:
        note_lines.append("[Note] Breame not installed: UK spelling conversion skipped. pip install breame")
    if _get_anthropic_client() is None:
        note_lines.append("[Note] Anthropic not configured: set ANTHROPIC_API_KEY in .env or 'env (1)'")

    if note_lines:
        preview = "\n".join(note_lines) + ("\n\n" + (preview or ""))

    return out_path, preview

with gr.Blocks(title="Style Rules Transformer + Claude (MAJOR_PROMPT)") as demo:
    gr.Markdown(
        "### Style Rules Transformer (Claude post-edit mandatory)\n"
        "- Upload **.docx** or text/markdown/HTML.\n"
        "- Deterministic rules first (numbers/percentages/ranges/currency/temperatures/UK spellings).\n"
        "- Then **Claude (Anthropic)** runs with your **MAJOR_PROMPT**. Quotes, tables, figures, headers, images, code fences, `<<par>>`, and asterisk-marked spans are preserved.\n"
        "- For **.docx**: Only body paragraphs (not tables/headings/quote styles) are edited; paragraph formatting is preserved."
    )
    with gr.Row():
        inp = gr.File(label="Upload .docx or text/markdown/HTML", file_count="single", type="filepath")
    with gr.Row():
        model = gr.Dropdown(
            choices=[
                "claude-3-5-sonnet-latest",
                "claude-3-5-haiku-latest"
            ],
            value="claude-3-5-sonnet-latest",
            label="Anthropic model"
        )
        max_tokens = gr.Slider(512, 8192, value=2048, step=64, label="Claude max tokens")
    with gr.Row():
        btn = gr.Button("Process")
    with gr.Row():
        out_file = gr.File(label="Download processed .docx")
    with gr.Row():
        preview = gr.Textbox(label="Preview (first ~80 lines & notes)", lines=20)

    btn.click(process_file, inputs=[inp, model, max_tokens], outputs=[out_file, preview])

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860, share=False)
