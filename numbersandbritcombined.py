#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import io
import json
import time
import tempfile
from typing import Dict, Tuple, List, Optional

import gradio as gr
from dotenv import load_dotenv
from docx import Document

# UK spelling helpers
from breame.spelling import (
    get_british_spelling,
    american_spelling_exists,
    british_spelling_exists,
)

# 🔗 bring in the external prompt
from isas_prompt import ISAS_PROMPT

# ============ Anthropic setup ============
load_dotenv("env (1)")
ANTHROPIC_MODEL = "claude-3-5-haiku-latest"

# ----------------------------------------
# Utility
# ----------------------------------------
def normalize_thin_spaces(s: str) -> str:
    return s.replace("\u2009", " ").replace("\u00a0", " ")

# ----------------------------------------
# Currencies (incl. FX to SGD via LLM)
# ----------------------------------------
CURRENCY_MAP = {
    "S$": ("SGD", "S$"), "SGD": ("SGD", "S$"),
    "US$": ("USD", "US$"), "USD": ("USD", "US$"), "$": ("USD", "US$"),
    "£": ("GBP", "£"), "GBP": ("GBP", "£"),
    "€": ("EUR", "€"), "EUR": ("EUR", "€"),
    "¥": ("JPY", "¥"), "JPY": ("JPY", "¥"),
    "₹": ("INR", "₹"), "INR": ("INR", "₹"),
}
CUR_SYMS = "|".join(map(re.escape, sorted(CURRENCY_MAP.keys(), key=len, reverse=True)))
SCALE_WORDS = {"million": 1_000_000, "billion": 1_000_000_000}

CURRENCY_RE = re.compile(
    rf"""
    (?P<cur>{CUR_SYMS})
    \s*
    (?P<amount>\d{{1,3}}(?:,\d{{3}})*|\d+)(?P<dec>\.\d+)?   # 12,345.67 or 123.45
    (?:\s*(?P<scale>million|billion))?                     # optional scale
    """, re.IGNORECASE | re.VERBOSE
)

def format_amount(symbol: str, amount_str: str, dec: Optional[str], scale: Optional[str]) -> str:
    core = amount_str + (dec or "")
    return f"{symbol}{core} {scale.lower()}" if scale else f"{symbol}{core}"

def already_has_sgd_parenthetical(text: str, start_idx: int) -> bool:
    tail = text[start_idx:start_idx+80]
    return bool(re.match(r"\s*\(S\$\s?\d", tail))

def normalize_currencies(text: str, rates_to_sgd: Dict[str, float]) -> str:
    def add_sgd_conversion(m: re.Match) -> str:
        cur = m.group("cur"); amount = m.group("amount"); dec = m.group("dec") or ""
        scale = (m.group("scale") or "").lower()
        code, symbol = CURRENCY_MAP[cur]
        src_disp = format_amount(symbol, amount, dec, scale)

        if code == "SGD" or already_has_sgd_parenthetical(text, m.end()):
            return src_disp

        amt = float(amount.replace(",", "")) + (float(dec) if dec else 0.0)
        multiplier = SCALE_WORDS.get(scale, 1)
        rate = rates_to_sgd.get(code)
        if not rate:
            return src_disp

        sgd_value = amt * multiplier * rate
        if scale:
            scaled = sgd_value / multiplier
            sgd_disp = f"S${scaled:,.2f} {scale}"
        else:
            sgd_disp = f"S${sgd_value:,.2f}"
        return f"{src_disp} ({sgd_disp})"

    out = CURRENCY_RE.sub(add_sgd_conversion, text)
    out = re.sub(r"(S\$|US\$|£|€|¥|₹)\s+(\d)", r"\1\2", out)
    out = re.sub(r"(\d(?:\.\d+)?)\s*(million|billion)\b", r"\1 \2", out, flags=re.IGNORECASE)
    return out

_fx_cache: Dict[str, float] = {}
_fx_cached_at: float = 0.0

def get_fx_rates_to_sgd(codes: List[str]) -> Dict[str, float]:
    """Get FX rates to SGD via Claude. Cache for 10 min."""
    import anthropic
    global _fx_cache, _fx_cached_at
    now = time.time()
    need = [c.upper() for c in codes if c.upper() != "SGD"]
    if not need:
        return {}
    if _fx_cache and (now - _fx_cached_at < 600):
        return {c: _fx_cache.get(c) for c in need if _fx_cache.get(c)}

    payload = ", ".join(sorted(set(need)))
    prompt = (
        "You are a precise FX assistant. Respond with STRICT JSON only, no prose. "
        "Provide latest reasonable spot exchange rates as of today for these currency codes, "
        "quoted as '1 UNIT = X SGD'. Return a JSON object with keys as codes and values as numbers.\n"
        f"Codes: [{payload}]\n"
        "Example output:\n"
        "{ \"USD\": 1.35, \"EUR\": 1.46 }\n"
        "Do not include comments, units, dates, or extra fields."
    )

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=400,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )
    text = "".join(getattr(b, "text", "") for b in msg.content)
    try:
        data = json.loads(text)
        rates = {k.upper(): float(v) for k, v in data.items()}
    except Exception:
        pairs = re.findall(r'"([A-Za-z]{3})"\s*:\s*([0-9]+(?:\.[0-9]+)?)', text)
        rates = {k.upper(): float(v) for k, v in pairs}

    for k, v in rates.items():
        if v and v > 0:
            _fx_cache[k] = v
    _fx_cached_at = time.time()
    return {c: _fx_cache.get(c) for c in need if _fx_cache.get(c)}

# ----------------------------------------
# Temperatures
# ----------------------------------------
def _degree_word(n: float) -> str:
    return "degree" if abs(n) == 1 else "degrees"

def normalize_temperatures(text: str) -> str:
    def repl_symbol(m: re.Match) -> str:
        num = float(m.group("num"))
        unit = m.group("unit").upper()
        unit_full = "Celsius" if unit == "C" else "Fahrenheit"
        return f"{m.group('rawnum')} {_degree_word(num)} {unit_full}"

    def repl_wordy(m: re.Match) -> str:
        num = float(m.group("num"))
        unit = m.group("unit").upper()
        unit_full = "Celsius" if unit == "C" else "Fahrenheit"
        return f"{m.group('rawnum')} {_degree_word(num)} {unit_full}"

    text = re.sub(r"(?P<rawnum>-?\d+(?:\.\d+)?)\s*°\s*(?P<unit>[cCfF])", repl_symbol, text)
    text = re.sub(r"(?P<rawnum>-?\d+(?:\.\d+)?)\s*(?:deg(?:ree)?s?)\s*(?P<unit>[cCfF])\b",
                  repl_wordy, text, flags=re.IGNORECASE)
    return text

# ----------------------------------------
# Numbers & percentages
# ----------------------------------------
NUM_WORDS = {
    0:"zero",1:"one",2:"two",3:"three",4:"four",5:"five",
    6:"six",7:"seven",8:"eight",9:"nine",10:"ten",11:"eleven",
    12:"twelve",13:"thirteen",14:"fourteen",15:"fifteen",
    16:"sixteen",17:"seventeen",18:"eighteen",19:"nineteen"
}
TENS = {20:"twenty",30:"thirty",40:"forty",50:"fifty",60:"sixty",70:"seventy",80:"eighty",90:"ninety"}

def int_to_words(n: int) -> str:
    if n < 0 or n > 999: return str(n)
    if n < 20: return NUM_WORDS[n]
    if n < 100:
        t, r = divmod(n, 10)
        return TENS[t*10] + (f"-{NUM_WORDS[r]}" if r else "")
    h, r = divmod(n, 100)
    if r == 0: return f"{NUM_WORDS[h]} hundred"
    if r < 20: return f"{NUM_WORDS[h]} hundred and {NUM_WORDS[r]}"
    t, o = divmod(r, 10)
    tail = TENS[t*10] + (f"-{NUM_WORDS[o]}" if o else "")
    return f"{NUM_WORDS[h]} hundred and {tail}"

def spell_out_single_digits(text: str) -> str:
    YEAR_RE = r"(?:(?:19|20)\d{2})"
    pattern = re.compile(
        rf"""
        (?<![\w-])
        (?!{YEAR_RE}\b)
        ([0-9])
        (?!\d)
        (?!st\b|nd\b|rd\b|th\b)
        (?!\s*[\u2013\u2014\-]\s*\d)
        (?!\.\d)
        (?![A-Za-z])
        """, re.VERBOSE
    )
    return pattern.sub(lambda m: NUM_WORDS[int(m.group())], text)

def _fix_sentence_start(sentence: str) -> str:
    s = sentence.lstrip()
    leading_ws = sentence[:len(sentence)-len(s)]
    m = re.match(r"^(\d{1,3})(\b|$)", s)
    if m:
        s = int_to_words(int(m.group(1))) + s[m.end():]
    return leading_ws + s

def no_sentence_starts_with_digit(text: str) -> str:
    parts = re.split(r"(\s*[.!?]\s+)", text)
    if len(parts) == 1:
        return _fix_sentence_start(parts[0])
    fixed = []
    for i in range(0, len(parts), 2):
        sent = parts[i]
        delim = parts[i+1] if i+1 < len(parts) else ""
        fixed.append(_fix_sentence_start(sent) + delim)
    return "".join(fixed)

def pluralize_units_for_decimals(text: str) -> str:
    def repl(m):
        num = float(m.group("num"))
        unit = m.group("unit")
        if num > 1.0 and not unit.lower().endswith("s"):
            return f"{m.group('prefix')}{m.group('num')} {unit}s"
        return m.group(0)
    return re.sub(
        r"(?P<prefix>\b)(?P<num>\d+\.\d+)\s+(?P<unit>[A-Za-z]+)\b",
        repl,
        text
    )

# Percentages (outside quotes only)
_WORD_NUMS_0_19 = {
    "zero":0,"one":1,"two":2,"three":3,"four":4,"five":5,"six":6,"seven":7,"eight":8,"nine":9,
    "ten":10,"eleven":11,"twelve":12,"thirteen":13,"fourteen":14,"fifteen":15,"sixteen":16,
    "seventeen":17,"eighteen":18,"nineteen":19
}
_WORD_TENS = {
    "twenty":20,"thirty":30,"forty":40,"fifty":50,"sixty":60,"seventy":70,"eighty":80,"ninety":90
}
def _words_to_int_0_99(s: str) -> Optional[int]:
    s = s.strip().lower()
    if s in _WORD_NUMS_0_19: return _WORD_NUMS_0_19[s]
    if s in _WORD_TENS: return _WORD_TENS[s]
    if "-" in s:
        a, b = s.split("-", 1)
        if a in _WORD_TENS and b in _WORD_NUMS_0_19:
            return _WORD_TENS[a] + _WORD_NUMS_0_19[b]
    return None

_QUOTE_SPLIT_RE = re.compile(r'(".*?"|“.*?”|‘.*?’|\'.*?\')', flags=re.DOTALL)
def _apply_outside_quotes(text: str, transform) -> str:
    parts = _QUOTE_SPLIT_RE.split(text)
    for i in range(len(parts)):
        if i % 2 == 0:
            parts[i] = transform(parts[i])
    return "".join(parts)

def normalize_percentages(text: str) -> str:
    def _tx(seg: str) -> str:
        if not seg: return seg
        seg = re.sub(
            r'(?P<a>\d+(?:\.\d+)?)\s*[-–]\s*(?P<b>\d+(?:\.\d+)?)\s*(%|percent\b|per\s*cent\b)',
            lambda m: f"{m.group('a')}–{m.group('b')} per cent",
            seg, flags=re.IGNORECASE
        )
        seg = re.sub(r'(?P<n>\d+(?:\.\d+)?)\s*%', lambda m: f"{m.group('n')} per cent", seg)
        seg = re.sub(r'(?P<n>\d+(?:\.\d+)?)\s*percent\b', lambda m: f"{m.group('n')} per cent", seg, flags=re.IGNORECASE)
        def repl_wordnum(m):
            val = _words_to_int_0_99(m.group('w'))
            return f"{val} per cent" if val is not None else m.group(0)
        seg = re.sub(r'\b(?P<w>[A-Za-z]+(?:-[A-Za-z]+)*)\s+(?:percent\b|per\s*cent\b)', repl_wordnum, seg, flags=re.IGNORECASE)
        seg = re.sub(r'\bpercent\b', 'per cent', seg, flags=re.IGNORECASE)
        return seg
    return _apply_outside_quotes(text, _tx)

# ----------------------------------------
# UK English (breame) + exceptions
# ----------------------------------------
ORG_EXCEPTIONS = {
    "World Health Organization",
    "World Trade Organization",
    "International Civil Aviation Organization",
}
ORG_EXCEPTIONS_LOWER = {s.lower(): s for s in ORG_EXCEPTIONS}
WORD_RE = re.compile(r"[A-Za-z][A-Za-z'\-]*")

def _is_acronym(token: str) -> bool:
    if len(token) <= 1: return False
    caps = sum(1 for c in token if c.isupper())
    return (caps / len(token)) >= 0.8

def _is_exception_window(text: str, start: int, end: int) -> bool:
    window = text[max(0, start-60):min(len(text), end+60)].lower()
    for exc in ORG_EXCEPTIONS_LOWER:
        if exc in window:
            return True
    return False

def _to_uk_spelling_token(token: str) -> str:
    if _is_acronym(token):
        return token
    base = token
    caps = base.isupper()
    title = base.istitle()
    if british_spelling_exists(base.lower()):
        return base
    if american_spelling_exists(base.lower()):
        uk = get_british_spelling(base.lower()) or base
        if caps: uk = uk.upper()
        elif title: uk = uk.capitalize()
        return uk
    return base

def britishise_text_with_exceptions(text: str) -> str:
    out = []
    idx = 0
    for m in WORD_RE.finditer(text):
        start, end = m.span()
        out.append(text[idx:start])
        token = m.group(0)
        if _is_exception_window(text, start, end):
            out.append(token)
        else:
            out.append(_to_uk_spelling_token(token))
        idx = end
    out.append(text[idx:])
    result = "".join(out)
    low = result.lower()
    for exc_low, exc_canon in ORG_EXCEPTIONS_LOWER.items():
        pos = 0
        while True:
            i = low.find(exc_low, pos)
            if i == -1:
                break
            result = result[:i] + exc_canon + result[i+len(exc_canon):]
            low = result.lower()
            pos = i + len(exc_canon)
    return result

CONSISTENT_CHOICES = {
    "Ulema": {"Uluma", "Ulama", "Ulamaa"},
    "Shia": {"Shiah", "Shi'a", "Shi‘a", "Shiite"},
}
def enforce_consistent_foreign_spellings(text: str) -> str:
    for canonical, variants in CONSISTENT_CHOICES.items():
        for var in sorted(variants, key=len, reverse=True):
            def repl(m):
                s = m.group(0)
                if s.isupper(): return canonical.upper()
                if s.istitle(): return canonical
                if s.islower(): return canonical.lower()
                return canonical
            text = re.sub(rf"\b{re.escape(var)}\b", repl, text, flags=re.IGNORECASE)
    return text

def apply_uk_language_rules(text: str) -> str:
    text = britishise_text_with_exceptions(text)
    text = enforce_consistent_foreign_spellings(text)
    return text

# ----------------------------------------
# Percentages + numbers + currencies + temps pipeline
# ----------------------------------------
def tag_currencies(text: str) -> Tuple[str, Dict[str, str]]:
    placeholders = {}
    def repl(m):
        key = f"§CUR{len(placeholders)}§"
        placeholders[key] = m.group(0)
        return key
    return CURRENCY_RE.sub(repl, text), placeholders

def untag_currencies(text: str, placeholders: Dict[str, str]) -> str:
    for k, v in placeholders.items():
        text = text.replace(k, v)
    return text

def apply_numbers_currencies_temps(text: str, rates_to_sgd: Dict[str, float]) -> str:
    text = normalize_thin_spaces(text)
    text = normalize_currencies(text, rates_to_sgd or {})
    protected, placeholders = tag_currencies(text)
    protected = normalize_temperatures(protected)
    protected = spell_out_single_digits(protected)
    protected = normalize_percentages(protected)     # ← percentages rule integrated
    protected = no_sentence_starts_with_digit(protected)
    protected = pluralize_units_for_decimals(protected)
    return untag_currencies(protected, placeholders)

# ----------------------------------------
# LLM foreign-term glossing (now with external system prompt)
# ----------------------------------------
def _batch_paragraphs_for_gloss(paragraphs: List[str], max_chars: int = 8000) -> List[List[Tuple[int, str]]]:
    chunks: List[List[Tuple[int, str]]] = []
    cur: List[Tuple[int, str]] = []
    cur_len = 0
    for idx, txt in enumerate(paragraphs):
        t = txt or ""
        tlen = len(t) + 50
        if cur and cur_len + tlen > max_chars:
            chunks.append(cur)
            cur = []
            cur_len = 0
        cur.append((idx, t))
        cur_len += tlen
    if cur:
        chunks.append(cur)
    return chunks

def _llm_gloss_batch(candidates: List[Tuple[int, str]]) -> Dict[int, List[Dict[str, str]]]:
    import anthropic
    if not candidates:
        return {}
    payload = [{"id": idx, "text": txt} for idx, txt in candidates]

    # The user prompt describing the JSON we want:
    user_prompt = (
        "For each paragraph item below, identify up to 6 clearly foreign (non-English) words or short phrases "
        "that should be italicised and provide a concise English gloss (≤ 5 words). "
        "Include romanised phrases like 'Viksit Bharat', 'Atmanirbhar Bharat', 'Swachh Bharat Abhiyan'.\n"
        "STRICTLY DO NOT include English words/phrases, English proper names, or acronyms (WHO, WTO, GDP, etc.).\n"
        "Return STRICT JSON only as:\n"
        "{ \"results\": [ {\"id\": int, \"items\": [ {\"phrase\": str, \"gloss\": str}, ...]}, ...] }\n"
        f"Items:\n{json.dumps(payload, ensure_ascii=False)}"
    )

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=1500,
        temperature=0,
        messages=[
            {"role": "system", "content": ISAS_PROMPT},   # ← use external prompt
            {"role": "user", "content": user_prompt},
        ],
    )
    text = "".join(getattr(b, "text", "") for b in msg.content)
    try:
        data = json.loads(text)
        results = data.get("results", [])
        out: Dict[int, List[Dict[str, str]]] = {}
        for entry in results:
            pid = entry.get("id")
            items = entry.get("items", [])
            if isinstance(pid, int) and isinstance(items, list):
                clean = []
                for it in items:
                    phrase = (it.get("phrase") or "").strip()
                    gloss = (it.get("gloss") or "").strip()
                    if phrase and gloss:
                        clean.append({"phrase": phrase, "gloss": gloss})
                out[pid] = clean[:6]
        return out
    except Exception:
        return {}

def italicise_and_gloss(paragraph, items: List[Dict[str, str]]):
    if not items:
        return
    items_sorted = sorted(items, key=lambda d: len(d["phrase"]), reverse=True)
    for it in items_sorted:
        phrase = it["phrase"]; gloss = it["gloss"]
        if not phrase or not gloss:
            continue
        for run in paragraph.runs:
            if not run.text:
                continue
            i = run.text.find(phrase)
            if i == -1:
                continue
            before = run.text[:i]
            after = run.text[i+len(phrase):]
            run.text = before
            phr_run = paragraph.add_run(phrase)
            phr_run.italic = True
            paragraph.add_run(f" ({gloss})")
            paragraph.add_run(after)

# ----------------------------------------
# DOCX processing (body only)
# ----------------------------------------
def process_document(input_bytes: bytes) -> bytes:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
        tmp_in.write(input_bytes)
        in_path = tmp_in.name

    doc = Document(in_path)

    # gather currencies to fetch FX once
    big_text = "\n".join(p.text for p in doc.paragraphs)
    currency_codes_found: List[str] = []
    for m in CURRENCY_RE.finditer(big_text):
        code = CURRENCY_MAP[m.group("cur")][0]
        if code not in currency_codes_found:
            currency_codes_found.append(code)
    rates_to_sgd = get_fx_rates_to_sgd(currency_codes_found) if currency_codes_found else {}

    # pass 1: numeric/currency/temp + UK rules
    for p in doc.paragraphs:
        txt = p.text
        txt = apply_numbers_currencies_temps(txt, rates_to_sgd)
        txt = apply_uk_language_rules(txt)
        if p.runs:
            p.runs[0].text = txt
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(txt)

    # pass 2: LLM foreign-term glossing (with external system prompt)
    paragraphs_text = [p.text for p in doc.paragraphs]
    chunks = _batch_paragraphs_for_gloss(paragraphs_text, max_chars=8000)
    phrase_cache: Dict[str, str] = {}
    for chunk in chunks:
        result_map = _llm_gloss_batch(chunk)
        for idx, items in result_map.items():
            dedup: List[Dict[str, str]] = []
            for it in items:
                ph = it["phrase"]
                gl = phrase_cache.get(ph) or it["gloss"]
                phrase_cache[ph] = gl
                dedup.append({"phrase": ph, "gloss": gl})
            italicise_and_gloss(doc.paragraphs[idx], dedup)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# ----------------------------------------
# Gradio app (minimal)
# ----------------------------------------
def convert_file(docx_path):
    if not docx_path:
        raise gr.Error("Please upload a .docx file.")
    with open(str(docx_path), "rb") as f:
        data = f.read()
    output_bytes = process_document(data)
    out_path = os.path.join(tempfile.gettempdir(), f"isas_style_{int(time.time())}.docx")
    with open(out_path, "wb") as f:
        f.write(output_bytes)
    return out_path

with gr.Blocks(title="ISAS Style — Numbers, Currencies, Temperatures & UK (body only)") as demo:
    gr.Markdown(
        "## ISAS Style — Numbers, Currencies, Temperatures & UK (Body paragraphs only)\n"
        "- Numbers: 1–9 spelled out; no sentence starts with a digit; decimals > 1 pluralise unit.\n"
        "- Currencies: add (S$...) for non-SGD via Claude 3.5 Haiku (cached).\n"
        "- Temperatures: e.g., 1.5°C → 1.5 degrees Celsius.\n"
        "- UK English: breame + exceptions; consistency for select terms.\n"
        "- Foreign terms: italicise + short English gloss (LLM; uses external prompt).\n\n"
        "**Note:** Tables, headers, footers, and notes are not processed."
    )
    inp = gr.File(label="Upload .docx", file_types=[".docx"], type="filepath")
    out = gr.File(label="Download processed .docx")
    go = gr.Button("Apply style")
    go.click(convert_file, inputs=[inp], outputs=out)

if __name__ == "__main__":
    demo.launch()
