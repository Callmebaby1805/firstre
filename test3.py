# import gradio as gr
# import json
# import re
# from collections import defaultdict
# from docx import Document

# # ---------------------------- Utility Functions ---------------------------- #

# def normalize_key(text):
#     return re.sub(r'\s+', '', text.lower())

# def load_abbr_map(json_path):
#     with open(json_path, "r", encoding="utf-8") as f:
#         return json.load(f)

# def clean_nested_expansion(text):
#     text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
#     text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
#     return text

# # ---------------------- Special Case: United States Fix -------------------- #

# def handle_special_case_usa(text):
#     pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
#     text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    
#     pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
#     text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)

#     return text

# # -------------------- Synonym Replacement (Global First) ------------------- #

# def resolve_synonyms_first_mention_global(text, synonym_groups):
#     for group in synonym_groups:
#         positions = {}
#         for term in group:
#             match = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
#             if match:
#                 positions[term] = match.start()

#         if positions:
#             chosen = min(positions, key=positions.get)
#             for alt in group:
#                 if alt.lower() != chosen.lower():
#                     pattern = rf'\b{re.escape(alt)}\b'
#                     text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)

#     return text

# # -------------------- Main Abbreviation Logic Processor -------------------- #

# def expand_abbreviations(doc_path, json_path="abbreviations.json"):
#     abbr_map = load_abbr_map(json_path)
#     counts = defaultdict(int)
#     norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

#     input_doc = Document(doc_path)
#     result_doc = Document()

#     synonym_groups = [
#         ["South Korea", "Republic of Korea", "ROK"],
#         ["The United States", "The United States of America"],
#         ["North Korea", "Democratic People's Republic of Korea", "DPRK"]
#     ]

#     # Utilities
#     def is_heading(paragraph):
#         return paragraph.style.name.startswith("Heading")

#     def is_in_table(paragraph):
#         return paragraph._element.getparent().tag.endswith('tbl')

#     header_texts = []
#     footer_texts = []
#     for section in input_doc.sections:
#         header_texts.extend(p.text.strip() for p in section.header.paragraphs)
#         footer_texts.extend(p.text.strip() for p in section.footer.paragraphs)

#     def is_in_header_or_footer(text):
#         return text.strip() in header_texts or text.strip() in footer_texts

#     # Step 1: Global Synonym Normalization (on body text only)
#     full_text = "\n".join([
#         p.text for p in input_doc.paragraphs
#         if not is_heading(p) and not is_in_table(p) and not is_in_header_or_footer(p.text)
#     ])
#     full_text = resolve_synonyms_first_mention_global(full_text, synonym_groups)

#     non_skipped_texts = full_text.split("\n")
#     idx = 0
#     for para in input_doc.paragraphs:
#         if is_heading(para) or is_in_table(para) or is_in_header_or_footer(para.text):
#             para.text = para.text  # untouched
#         else:
#             para.text = non_skipped_texts[idx]
#             idx += 1

#     # Step 2: Abbreviation Expansion
#     for para in input_doc.paragraphs:
#         line = para.text

#         if is_heading(para) or is_in_table(para) or is_in_header_or_footer(line):
#             result_doc.add_paragraph(line)  # skip processing
#             continue

#         line = clean_nested_expansion(line)

#         for norm_exp, (abbr, expansion) in norm_map.items():
#             abbr_pattern = rf'\b{re.escape(abbr)}\b'
#             exp_pattern = rf'\b{re.escape(expansion)}\b'
#             patterns = [
#                 rf'{abbr}\s*\(\s*{expansion}\s*\)',
#                 rf'{expansion}\s*\(\s*{abbr}\s*\)',
#                 rf'{abbr}\s*\(\s*\w[\w\s]*\)',
#                 rf'\w[\w\s]*\(\s*{abbr}\s*\)',
#                 abbr_pattern,
#                 exp_pattern
#             ]

#             if any(re.search(p, line, re.IGNORECASE) for p in patterns):
#                 counts[abbr] += 1
#                 if counts[abbr] == 1:
#                     correct_format = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
#                     if not re.search(correct_format, line, re.IGNORECASE):
#                         line = re.sub('|'.join(patterns), f"{expansion} ({abbr})", line, count=1, flags=re.IGNORECASE)
#                 else:
#                     line = re.sub('|'.join(patterns), abbr, line, flags=re.IGNORECASE)

#         line = handle_special_case_usa(line)
#         result_doc.add_paragraph(line)

#     out_path = "expanded_result.docx"
#     result_doc.save(out_path)
#     return out_path

# # ---------------------------- Gradio Interface ----------------------------- #

# def gradio_abbreviation_expand(doc_file, json_file):
#     output_path = expand_abbreviations(doc_file.name, json_file.name)
#     return "✅ Document processed successfully!", output_path

# gr.Interface(
#     fn=gradio_abbreviation_expand,
#     inputs=[
#         gr.File(label="Upload DOCX Document", file_types=[".docx"]),
#         gr.File(label="Upload Abbreviation JSON", file_types=[".json"])
#     ],
#     outputs=[
#         gr.Text(label="Status"),
#         gr.File(label="Download Processed File")
#     ],
#     title="Abbreviation Expander (Controlled)",
#     description="Expands abbreviations on first use as 'Expansion (ABBR)', later uses as 'ABBR'. Skips headers, footers, headings, and tables."
# ).launch()
import gradio as gr
import json
import re
from collections import defaultdict
from docx import Document

# ---------------------------- Utility Functions ---------------------------- #

def normalize_key(text):
    return re.sub(r'\s+', '', text.lower())

def load_abbr_map(json_path="abbreviations.json"):
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)

def clean_nested_expansion(text):
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

# ---------------------- Special Case: United States Fix -------------------- #

def handle_special_case_usa(text):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)

    return text

# -------------------- Synonym Replacement (Global First) ------------------- #

def resolve_synonyms_first_mention_global(text, synonym_groups):
    for group in synonym_groups:
        positions = {}
        for term in group:
            match = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
            if match:
                positions[term] = match.start()

        if positions:
            chosen = min(positions, key=positions.get)
            for alt in group:
                if alt.lower() != chosen.lower():
                    pattern = rf'\b{re.escape(alt)}\b'
                    text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)

    return text

# -------------------- Main Abbreviation Logic Processor -------------------- #

def expand_abbreviations(doc_path, json_path="abbreviations.json"):
    abbr_map = load_abbr_map(json_path)
    counts = defaultdict(int)
    norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

    input_doc = Document(doc_path)
    result_doc = Document()

    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China","People Republic of China", "PRC" ]
    ]

    # Utilities
    def is_heading(paragraph):
        return paragraph.style.name.startswith("Heading")

    def is_in_table(paragraph):
        return paragraph._element.getparent().tag.endswith('tbl')

    header_texts = []
    footer_texts = []
    for section in input_doc.sections:
        header_texts.extend(p.text.strip() for p in section.header.paragraphs)
        footer_texts.extend(p.text.strip() for p in section.footer.paragraphs)

    def is_in_header_or_footer(text):
        return text.strip() in header_texts or text.strip() in footer_texts

    # Step 1: Global Synonym Normalization (on body text only)
    full_text = "\n".join([
        p.text for p in input_doc.paragraphs
        if not is_heading(p) and not is_in_table(p) and not is_in_header_or_footer(p.text)
    ])
    full_text = resolve_synonyms_first_mention_global(full_text, synonym_groups)

    non_skipped_texts = full_text.split("\n")
    idx = 0
    for para in input_doc.paragraphs:
        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(para.text):
            para.text = para.text  # untouched
        else:
            para.text = non_skipped_texts[idx]
            idx += 1

    # Step 2: Abbreviation Expansion
    for para in input_doc.paragraphs:
        line = para.text

        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(line):
            result_doc.add_paragraph(line)  # skip processing
            continue

        line = clean_nested_expansion(line)

        for norm_exp, (abbr, expansion) in norm_map.items():
            abbr_pattern = rf'\b{re.escape(abbr)}\b'
            exp_pattern = rf'\b{re.escape(expansion)}\b'
            patterns = [
                rf'{abbr}\s*\(\s*{expansion}\s*\)',
                rf'{expansion}\s*\(\s*{abbr}\s*\)',
                rf'{abbr}\s*\(\s*\w[\w\s]*\)',
                rf'\w[\w\s]*\(\s*{abbr}\s*\)',
                abbr_pattern,
                exp_pattern
            ]

            if any(re.search(p, line, re.IGNORECASE) for p in patterns):
                counts[abbr] += 1
                if counts[abbr] == 1:
                    correct_format = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
                    if not re.search(correct_format, line, re.IGNORECASE):
                        line = re.sub('|'.join(patterns), f"{expansion} ({abbr})", line, count=1, flags=re.IGNORECASE)
                else:
                    line = re.sub('|'.join(patterns), abbr, line, flags=re.IGNORECASE)

        line = handle_special_case_usa(line)
        result_doc.add_paragraph(line)

    out_path = "expanded_result.docx"
    result_doc.save(out_path)
    return out_path

# ---------------------------- Gradio Interface ----------------------------- #

def gradio_abbreviation_expand(doc_file):
    output_path = expand_abbreviations(doc_file.name)
    return "✅ Document processed successfully!", output_path

gr.Interface(
    fn=gradio_abbreviation_expand,
    inputs=[
        gr.File(label="Upload DOCX Document", file_types=[".docx"])
    ],
    outputs=[
        gr.Text(label="Status"),
        gr.File(label="Download Processed File")
    ],
    title="Abbreviation Expander (Controlled)",
    description="Expands abbreviations on first use as 'Expansion (ABBR)', later uses as 'ABBR'. Loads local 'abbreviations.json'. Skips headers, footers, headings, and tables."
).launch()
