# #!/usr/bin/env python3
# # -*- coding: utf-8 -*-
# import re
# import time
# from pathlib import Path
# from typing import Optional, Tuple
# import gradio as gr

# BREAME = None
# try:
#     from breame.spelling import get_british_spelling  # type: ignore
#     BREAME = True
# except Exception:
#     BREAME = False

# QUOTE_RE = re.compile(r'(".*?"|\'.*?\'|“[^”]*”|‘[^’]*’)', re.S)
# STAR_RE  = re.compile(r'(\*\*.*?\*\*|\*.*?\*)', re.S)
# PAR_MARK = '<<par>>'
# WORD_RE = re.compile(r"\b([A-Za-z][A-Za-z'-]*)\b")

# US_UK = {
#     "color": "colour", "colors": "colours", "colored": "coloured", "coloring": "colouring",
#     "favorite": "favourite", "favorites": "favourites",
#     "behavior": "behaviour", "behaviors": "behaviours",
#     "honor": "honour", "honors": "honours", "honorable": "honourable",
#     "labor": "labour", "neighbor": "neighbour", "neighbors": "neighbours",
#     "vigor": "vigour",
#     "center": "centre", "centers": "centres",
#     "meter": "metre", "meters": "metres",
#     "liter": "litre", "liters": "litres",
#     "fiber": "fibre",
#     "defense": "defence", "pretense": "pretence",
#     "organize": "organise", "organizes": "organises", "organizing": "organising", "organized": "organised",
#     "organization": "organisation", "organizations": "organisations",
#     "realize": "realise", "realized": "realised", "realizes": "realises", "realizing": "realising",
#     "realization": "realisation",
#     "traveling": "travelling", "traveled": "travelled", "traveler": "traveller", "travelers": "travellers",
#     "catalog": "catalogue", "dialog": "dialogue",
#     "jewelry": "jewellery",
#     "aluminum": "aluminium",
#     "tire": "tyre",
#     "curb": "kerb",
# }

# ISE_EXCEPTIONS = {"size", "seize", "capsize", "prize"}

# def case_preserve(repl: str, src: str) -> str:
#     if src.isupper():
#         return repl.upper()
#     if src[0].isupper():
#         return repl[0].upper() + repl[1:]
#     return repl

# def britishize_tokens(text: str) -> str:
#     out = []
#     last = 0
#     for m in WORD_RE.finditer(text):
#         word = m.group(1)
#         low = word.lower()
#         repl = None
#         if BREAME:
#             try:
#                 from breame.spelling import get_british_spelling
#                 tmp = get_british_spelling(low)
#                 if tmp:
#                     repl = case_preserve(tmp, word)
#             except Exception:
#                 repl = None
#         if repl is None:
#             import re as _re
#             if low in US_UK:
#                 repl = case_preserve(US_UK[low], word)
#             elif _re.search(r"ization\b", low):
#                 repl = _re.sub(r"ization\b", "isation", word, flags=_re.I)
#             elif _re.search(r"ize\b", low) and low not in ISE_EXCEPTIONS:
#                 repl = _re.sub(r"ize\b", "ise", word, flags=_re.I)
#             else:
#                 repl = word
#         out.append(text[last:m.start()]); out.append(repl); last = m.end()
#     out.append(text[last:])
#     return "".join(out)

# def mask_patterns(text: str, pattern, start_index: int, mapping: dict) -> str:
#     def repl(m):
#         idx = len(mapping) + start_index
#         ph = f"\\uE000{idx}\\uE000"
#         mapping[ph] = m.group(0)
#         return ph
#     return pattern.sub(repl, text)

# def unmask(text: str, mapping: dict) -> str:
#     for ph, original in mapping.items():
#         text = text.replace(ph, original)
#     return text

# def britishize_text(text: str) -> str:
#     mapping = {}
#     masked = mask_patterns(text, STAR_RE, 0, mapping)
#     masked = mask_patterns(masked, QUOTE_RE, len(mapping), mapping)
#     idx = len(mapping)
#     while PAR_MARK in masked:
#         ph = f"\\uE000{idx}\\uE000"; mapping[ph] = PAR_MARK
#         masked = masked.replace(PAR_MARK, ph, 1); idx += 1
#     converted = britishize_tokens(masked)
#     return unmask(converted, mapping)

# def paragraph_is_in_table(paragraph) -> bool:
#     elt = paragraph._element
#     while elt is not None:
#         if elt.tag.endswith('}tc'): return True
#         elt = elt.getparent()
#     return False

# def is_footnote_paragraph(paragraph) -> bool:
#     elt = paragraph._element
#     while elt is not None:
#         tag = elt.tag.lower()
#         if tag.endswith('}footnote') or tag.endswith('}endnote'): return True
#         elt = elt.getparent()
#     return False

# def process_docx(input_path: str, output_path: str, process_tables: bool = False):
#     try:
#         from docx import Document
#     except ImportError:
#         return False, "python-docx is required. Install with: pip install python-docx"
#     try:
#         doc = Document(input_path)
#         for p in doc.paragraphs:
#             if is_footnote_paragraph(p): continue
#             if (not process_tables) and paragraph_is_in_table(p): continue
#             for r in p.runs:
#                 if r.text: r.text = britishize_text(r.text)
#         doc.save(output_path)
#         return True, None
#     except Exception as e:
#         return False, f"Error processing DOCX: {e}"

# def process_txt(input_path: str, output_path: str):
#     try:
#         data = Path(input_path).read_text(encoding='utf-8')
#         Path(output_path).write_text(britishize_text(data), encoding='utf-8')
#         return True, None
#     except Exception as e:
#         return False, f"Error processing TXT: {e}"

# def convert_file(file, process_tables=False):
#     if file is None:
#         return None, "Please upload a .docx or .txt file.", ""
#     start = time.time()
#     in_path = Path(file.name)
#     suffix = in_path.suffix.lower()
#     out_path = Path.cwd() / f"{in_path.stem}_UK{suffix}"
#     if suffix == ".docx":
#         ok, err = process_docx(str(in_path), str(out_path), process_tables=process_tables)
#     elif suffix == ".txt":
#         ok, err = process_txt(str(in_path), str(out_path))
#     else:
#         return None, "Only .docx and .txt are supported.", ""
#     duration = time.time() - start
#     if not ok: return None, f"❌ {err}", f"Processing time: {duration:.2f} s"
#     preview = Path(out_path).read_text(encoding='utf-8')[:2000] if suffix == ".txt" else "DOCX processed. Footnotes/endnotes were not modified. Download to review."
#     return str(out_path), preview, f"✅ Done in {duration:.2f} s (breame={'on' if BREAME else 'off'})"

# with gr.Blocks(title="US → UK Converter") as demo:
#     gr.Markdown("## US → UK British English Converter\nConverts .docx or .txt while preserving quotes, *stars*, and `<<par>>`.\n\n**Footnotes & endnotes are never modified.** Tables are skipped by default for DOCX.")
#     with gr.Row():
#         inp = gr.File(label="Upload .docx or .txt", file_types=[".docx", ".txt"])
#         tables = gr.Checkbox(label="Process tables (DOCX)", value=False)
#     btn = gr.Button("Convert to British English", variant="primary")
#     with gr.Row():
#         out_file = gr.File(label="Converted file")
#     with gr.Row():
#         preview = gr.Textbox(label="Preview (for .txt)", lines=12)
#     runtime = gr.Markdown()
#     btn.click(convert_file, inputs=[inp, tables], outputs=[out_file, preview, runtime])

# if __name__ == "__main__":
#     demo.launch()
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import re
import time
from pathlib import Path
from typing import List, Tuple, Optional, Callable
import gradio as gr

# ======================
# Optional: breame usage
# ======================
BREAME = None
try:
    from breame.spelling import get_british_spelling  # type: ignore
    BREAME = True
except Exception:
    BREAME = False

# ======================
# Shared regex & markers
# ======================
QUOTE_RE = re.compile(r'(".*?"|\'.*?\'|“[^”]*”|‘[^’]*’)', re.S)
STAR_RE  = re.compile(r'(\*\*.*?\*\*|\*.*?\*)', re.S)
PAR_MARK = '<<par>>'
WORD_RE  = re.compile(r"\b([A-Za-z][A-Za-z'-]*)\b")

# Footnote placeholder (from your LLM/processing pipeline)
FOOTNOTE_PAT = re.compile(r"§FOOTNOTE:\d+§")

# ====================================
# US→UK spellings (fallback dictionary)
# ====================================
US_UK = {
    "color": "colour", "colors": "colours", "colored": "coloured", "coloring": "colouring",
    "favorite": "favourite", "favorites": "favourites",
    "behavior": "behaviour", "behaviors": "behaviours",
    "honor": "honour", "honors": "honours", "honorable": "honourable",
    "labor": "labour", "neighbor": "neighbour", "neighbors": "neighbours",
    "vigor": "vigour",
    "center": "centre", "centers": "centres",
    "meter": "metre", "meters": "metres",
    "liter": "litre", "liters": "litres",
    "fiber": "fibre",
    "defense": "defence", "pretense": "pretence",
    "organize": "organise", "organizes": "organises", "organizing": "organising", "organized": "organised",
    "organization": "organisation", "organizations": "organisations",
    "realize": "realise", "realized": "realised", "realizes": "realises", "realizing": "realising",
    "realization": "realisation",
    "traveling": "travelling", "traveled": "travelled", "traveler": "traveller", "travelers": "travellers",
    "catalog": "catalogue", "dialog": "dialogue",
    "jewelry": "jewellery",
    "aluminum": "aluminium",
    "tire": "tyre",
    "curb": "kerb",
}
ISE_EXCEPTIONS = {"size", "seize", "capsize", "prize"}

def case_preserve(repl: str, src: str) -> str:
    if src.isupper():
        return repl.upper()
    if src[:1].isupper():
        return repl[:1].upper() + repl[1:]
    return repl

# ======================================
# Masking helpers (quotes, stars, <<par>>)
# ======================================
def mask_patterns(text: str, pattern, start_index: int, mapping: dict) -> str:
    def repl(m):
        idx = len(mapping) + start_index
        ph = f"\\uE000{idx}\\uE000"
        mapping[ph] = m.group(0)
        return ph
    return pattern.sub(repl, text)

def unmask(text: str, mapping: dict) -> str:
    for ph, original in mapping.items():
        text = text.replace(ph, original)
    return text

# =============================
# Britishise tokens (your code)
# =============================
def britishize_tokens(text: str) -> str:
    out = []
    last = 0
    for m in WORD_RE.finditer(text):
        word = m.group(1)
        low = word.lower()
        repl = None
        if BREAME:
            try:
                tmp = get_british_spelling(low)  # type: ignore
                if tmp:
                    repl = case_preserve(tmp, word)
            except Exception:
                repl = None
        if repl is None:
            if low in US_UK:
                repl = case_preserve(US_UK[low], word)
            elif re.search(r"ization\b", low):
                repl = re.sub(r"ization\b", "isation", word, flags=re.I)
            elif re.search(r"ize\b", low) and low not in ISE_EXCEPTIONS:
                repl = re.sub(r"ize\b", "ise", word, flags=re.I)
            else:
                repl = word
        out.append(text[last:m.start()]); out.append(repl); last = m.end()
    out.append(text[last:])
    return "".join(out)

def britishize_text(text: str) -> str:
    mapping = {}
    masked = mask_patterns(text, STAR_RE, 0, mapping)
    masked = mask_patterns(masked, QUOTE_RE, len(mapping), mapping)
    idx = len(mapping)
    while PAR_MARK in masked:
        ph = f"\\uE000{idx}\\uE000"; mapping[ph] = PAR_MARK
        masked = masked.replace(PAR_MARK, ph, 1); idx += 1
    converted = britishize_tokens(masked)
    return unmask(converted, mapping)

# =====================================================
# Rules 1–5 (Numbers, Fractions, Currency – no SGD now)
# =====================================================
ITALIC_ONE_PAT = re.compile(r"\*[^*\n]+\*")
ITALIC_TWO_PAT = re.compile(r"\*\*[^*\n]+\*\*")

def quote_spans(text: str) -> List[Tuple[int, int]]:
    spans = []
    stack = []
    for i, ch in enumerate(text):
        if ch in ("'", '"'):
            if stack and stack[-1][0] == ch:
                _, start = stack.pop()
                spans.append((start, i + 1))
            else:
                stack.append((ch, i))
    return spans

def protected_spans(text: str) -> List[Tuple[int, int]]:
    spans = quote_spans(text)
    for pat in (FOOTNOTE_PAT, ITALIC_TWO_PAT, ITALIC_ONE_PAT):
        for m in pat.finditer(text):
            spans.append((m.start(), m.end()))
    spans.sort()
    merged = []
    for s, e in spans:
        if not merged or s > merged[-1][1]:
            merged.append((s, e))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
    return merged

def replace_outside(text: str, regex: re.Pattern, repl: Callable[[re.Match], str]) -> str:
    spans = protected_spans(text)
    out = []
    idx = 0
    for m in regex.finditer(text):
        if any(a <= m.start() < b for a, b in spans):
            continue
        out.append(text[idx:m.start()])
        out.append(repl(m))
        idx = m.end()
    out.append(text[idx:])
    return ''.join(out)

# (1) Spell out single digits (1–9)
SINGLE_DIGITS = {
    "1": "one", "2": "two", "3": "three", "4": "four", "5": "five",
    "6": "six", "7": "seven", "8": "eight", "9": "nine"
}
SINGLE_DIGIT_PAT = re.compile(r"\b([1-9])\b")

def rule1_spellout_single_digits(text: str) -> str:
    def repl(m: re.Match) -> str:
        return SINGLE_DIGITS[m.group(1)]
    return replace_outside(text, SINGLE_DIGIT_PAT, repl)

# Small number→words (sentence starts)
ONES = ["zero","one","two","three","four","five","six","seven","eight","nine"]
TEENS = ["ten","eleven","twelve","thirteen","fourteen","fifteen","sixteen","seventeen","eighteen","nineteen"]
TENS = ["","","twenty","thirty","forty","fifty","sixty","seventy","eighty","ninety"]

def num_to_words_0_999(n: int) -> str:
    assert 0 <= n <= 999
    if n < 10: return ONES[n]
    if n < 20: return TEENS[n-10]
    if n < 100:
        t, r = divmod(n, 10)
        return TENS[t] if r == 0 else f"{TENS[t]} {ONES[r]}"
    h, r = divmod(n, 100)
    if r == 0: return f"{ONES[h]} hundred"
    return f"{ONES[h]} hundred {num_to_words_0_999(r)}"

def num_to_words(n: int) -> str:
    if n < 0: return "minus " + num_to_words(-n)
    if n <= 999: return num_to_words_0_999(n)
    if n <= 999_999:
        k, r = divmod(n, 1000)
        if r == 0: return f"{num_to_words_0_999(k)} thousand"
        return f"{num_to_words_0_999(k)} thousand {num_to_words_0_999(r)}"
    return str(n)  # keep big numbers as numerals

# (2) Optional words (>=10) → numerals
WORD_NUM_MAP = {
    **{w: i for i, w in enumerate(ONES)},
    **{w: 10 + i for i, w in enumerate(TEENS)},
    **{TENS[i]: 10 * i for i in range(2, 10)},
    "hundred": 100, "thousand": 1000
}
WORD_NUMBER_PAT = re.compile(r"\b(zero|one|two|three|four|five|six|seven|eight|nine|"
                             r"ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|"
                             r"twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety"
                             r"(?:\s+hundred)?(?:\s+(?:one|two|three|four|five|six|seven|eight|nine))?"
                             r"(?:\s+thousand(?:\s+\w+)?)?)\b", re.IGNORECASE)

def words_to_int_safe(s: str) -> Optional[int]:
    parts = s.lower().split()
    total = 0
    current = 0
    for p in parts:
        if p not in WORD_NUM_MAP:
            return None
        val = WORD_NUM_MAP[p]
        if val == 100:
            current = (current or 1) * 100
        elif val == 1000:
            total += (current or 1) * 1000
            current = 0
        else:
            current += val
    return total + current

def rule2_words_to_numerals_10plus(text: str, enable: bool = True) -> str:
    if not enable:
        return text
    def repl(m: re.Match) -> str:
        tok = m.group(0)
        val = words_to_int_safe(tok)
        if val is None or val < 10:
            return tok
        return str(val)
    return replace_outside(text, WORD_NUMBER_PAT, repl)

# (3) Sentence-initial number → words (Capitalised)
SENTENCE_INITIAL_NUM_PAT = re.compile(r"(?m)(^|\n)(?P<prefix>\s*)(?P<num>\d{1,6})(\b)")

def rule3_sentence_initial_number_to_words(text: str) -> str:
    spans = protected_spans(text)
    out, idx = [], 0
    for m in SENTENCE_INITIAL_NUM_PAT.finditer(text):
        start = m.start("num")
        if any(a <= start < b for a, b in spans):
            continue
        out.append(text[idx:m.start("num")])
        n = int(m.group("num"))
        words = num_to_words(n)
        if words:
            words = words[0].upper() + words[1:]
        out.append(words)
        idx = m.end("num")
    out.append(text[idx:])
    return ''.join(out)

# (4) Fractions > 1 → plural verb "are" (narrow heuristic)
FRACTION_VERB_PAT = re.compile(r"\b(?P<num>\d+(?:\.\d+)?)\s+(?P<noun>[A-Za-z]+)\s+is\b")

def rule4_fraction_plural_verb(text: str) -> str:
    def repl(m: re.Match) -> str:
        try:
            val = float(m.group("num").replace(',', ''))
        except ValueError:
            return m.group(0)
        if val > 1.0:
            return f"{m.group('num')} {m.group('noun')} are"
        return m.group(0)
    return replace_outside(text, FRACTION_VERB_PAT, repl)

# (5) Currency: spacing normalisation only (NO conversion)
CURRENCY_PAT = re.compile(
    r"(?P<cur>US\$|S\$|£|€|¥|₹)\s?(?P<amt>\d{1,3}(?:,\d{3})*(?:\.\d+)?)"
    r"(?:\s*(?P<scale>million|billion))?\b"
)

def normalize_currency(cur: str, amt: str, scale: Optional[str]) -> str:
    amt_no_spaces = amt.replace(" ", "")
    return f"{cur}{amt_no_spaces} {scale}" if scale else f"{cur}{amt_no_spaces}"

def rule5_currency_spacing_only(text: str) -> str:
    spans = protected_spans(text)
    out, idx = [], 0
    for m in CURRENCY_PAT.finditer(text):
        if any(a <= m.start() < b for a, b in spans):
            continue
        out.append(text[idx:m.start()])
        cur = m.group("cur"); amt = m.group("amt"); scale = m.group("scale")
        out.append(normalize_currency(cur, amt, scale))  # no (S$...) appended
        idx = m.end()
    out.append(text[idx:])
    return ''.join(out)

def apply_rules_1_to_5(
    text: str,
    *,
    convert_words_10_plus_to_numerals: bool = True
) -> str:
    text = rule3_sentence_initial_number_to_words(text)
    text = rule1_spellout_single_digits(text)
    text = rule2_words_to_numerals_10plus(text, enable=convert_words_10_plus_to_numerals)
    text = rule4_fraction_plural_verb(text)
    text = rule5_currency_spacing_only(text)  # spacing only, no conversion
    return text

# ====================================
# Combined transform: Rules 1–5 + UK
# ====================================
def combined_transform(
    text: str,
    *,
    convert_words_10_plus_to_numerals: bool
) -> str:
    t1 = apply_rules_1_to_5(
        text,
        convert_words_10_plus_to_numerals=convert_words_10_plus_to_numerals
    )
    t2 = britishize_text(t1)
    return t2

# =======================
# DOCX helpers (safe I/O)
# =======================
def paragraph_is_in_table(paragraph) -> bool:
    elt = paragraph._element
    while elt is not None:
        if elt.tag.endswith('}tc'): return True
        elt = elt.getparent()
    return False

def is_footnote_paragraph(paragraph) -> bool:
    elt = paragraph._element
    while elt is not None:
        tag = elt.tag.lower()
        if tag.endswith('}footnote') or tag.endswith('}endnote'): return True
        elt = elt.getparent()
    return False

def process_docx(
    input_path: str,
    output_path: str,
    *,
    process_tables: bool,
    convert_words_10_plus_to_numerals: bool
):
    try:
        from docx import Document
    except ImportError:
        return False, "python-docx is required. Install with: pip install python-docx"
    try:
        doc = Document(input_path)
        for p in doc.paragraphs:
            if is_footnote_paragraph(p): 
                continue
            if (not process_tables) and paragraph_is_in_table(p):
                continue
            for r in p.runs:
                if r.text:
                    r.text = combined_transform(
                        r.text,
                        convert_words_10_plus_to_numerals=convert_words_10_plus_to_numerals
                    )
        doc.save(output_path)
        return True, None
    except Exception as e:
        return False, f"Error processing DOCX: {e}"

def process_txt(
    input_path: str,
    output_path: str,
    *,
    convert_words_10_plus_to_numerals: bool
):
    try:
        data = Path(input_path).read_text(encoding='utf-8')
        out = combined_transform(
            data,
            convert_words_10_plus_to_numerals=convert_words_10_plus_to_numerals
        )
        Path(output_path).write_text(out, encoding='utf-8')
        return True, None
    except Exception as e:
        return False, f"Error processing TXT: {e}"

# ===========
# Gradio app
# ===========
def convert_file(file, process_tables, words_to_numerals):
    if file is None:
        return None, "Please upload a .docx or .txt file.", ""
    start = time.time()
    in_path = Path(file.name)
    suffix = in_path.suffix.lower()
    out_path = Path.cwd() / f"{in_path.stem}_ISAS_UK{suffix}"
    if suffix == ".docx":
        ok, err = process_docx(
            str(in_path), str(out_path),
            process_tables=process_tables,
            convert_words_10_plus_to_numerals=words_to_numerals
        )
        preview = "DOCX processed. Footnotes/endnotes preserved. Download to review."
    elif suffix == ".txt":
        ok, err = process_txt(
            str(in_path), str(out_path),
            convert_words_10_plus_to_numerals=words_to_numerals
        )
        preview = Path(out_path).read_text(encoding='utf-8')[:2000] if ok else ""
    else:
        return None, "Only .docx and .txt are supported.", ""
    duration = time.time() - start
    if not ok:
        return None, f"❌ {err}", f"Processing time: {duration:.2f} s"
    return str(out_path), preview, f"✅ Done in {duration:.2f} s (breame={'on' if BREAME else 'off'})"

with gr.Blocks(title="ISAS Numbers/Currency (No SGD) + US→UK Converter") as demo:
    gr.Markdown(
        "## ISAS Numbers/Currency Rules (1–5, **no SGD conversion**) + US→UK British English Converter\n"
        "- Preserves quotes, *stars*, `<<par>>`, and footnote placeholders like `§FOOTNOTE:12§`\n"
        "- DOCX: skips footnotes/endnotes and (by default) tables\n"
        "- Currency: **spacing normalisation only**; conversions are left to your LLM\n"
    )
    with gr.Row():
        inp = gr.File(label="Upload .docx or .txt", file_types=[".docx", ".txt"])
    with gr.Row():
        tables = gr.Checkbox(label="Process tables (DOCX)", value=False)
        words_to_nums = gr.Checkbox(label="Convert words ≥10 to numerals", value=True)
    btn = gr.Button("Convert", variant="primary")
    with gr.Row():
        out_file = gr.File(label="Converted file")
    with gr.Row():
        preview = gr.Textbox(label="Preview (for .txt)", lines=12)
    runtime = gr.Markdown()
    btn.click(convert_file, inputs=[inp, tables, words_to_nums], outputs=[out_file, preview, runtime])

if __name__ == "__main__":
    demo.launch()
