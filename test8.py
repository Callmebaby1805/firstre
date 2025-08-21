# # test8.py — Abbreviation expander (run LAST in pipeline)
# Core logic preserved from your original:
# - Topic classification (Anthropic) → merge base + category maps
# - First mention => "Full (ABBR)"; subsequent mentions => "ABBR"
# - Quote-aware (do not modify inside quotes)
# - Skip headings, tables, headers/footers
# - Synonym pre-pass (your original groups)
# - Names: collapse to last name after first full-name mention
# - USA special case
# - Import-safe (no Gradio UI)
#
# Patches applied:
# - Robust JSON loading (current dir + /mnt/data) and direction normalisation (abbr↔full)
# - Regex anchors changed to lookarounds (?<!\w … (?!\w)) so acronyms like "U.S." match
# - Optional out_path argument for pipeline integration

import os
import json
import re
import anthropic
from collections import defaultdict
from docx import Document
from dotenv import load_dotenv

# Load environment variables (your key is in "env (1)")
load_dotenv("env (1)")

# Anthropic client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# Your original category map
CATEGORY_TO_JSON = {
    "Sustainability and the Environment": "Sustainabilityandenvironment.json",
    "Strategic Technologies": "StatergicTechnologies.json",  # original spelling kept
    "Trade and Economics": "Trade.json",
    "Politics, Society and Governance": "PSG.json",
    "International Relations, Multipolarity and Multilateralism": "IEMTMLT.json",
}

# ---------------- utilities ----------------
def _candidates(filename: str):
    return [os.path.join(os.getcwd(), filename), os.path.join("/mnt/data", filename)]

def _load_json_first_found(filename: str) -> dict:
    for p in _candidates(filename):
        if os.path.exists(p):
            try:
                with open(p, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    return data
            except Exception:
                pass
    return {}

def _load_any_abbrev_map(*candidate_names) -> dict:
    # tries each name in both CWD and /mnt/data
    names = list(candidate_names) or ["abbreviations.json", "abbrevations.json"]
    for name in names:
        m = _load_json_first_found(name)
        if m:
            return m
    return {}

def normalize_key(text: str) -> str:
    return re.sub(r"\s+", "", text.lower())

def _is_acronym(s: str) -> bool:
    s = s.strip()
    if not s:
        return False
    if " " in s:
        return False
    letters = re.sub(r"[^A-Za-z]", "", s)
    return len(letters) > 0 and letters.upper() == letters and len(s) <= 15

def _to_pairs(abbr_map_raw: dict):
    """
    Convert a raw dict (which may be abbr->full or full->abbr) into
    canonical list of (abbr, full) pairs. Heuristics used only if ambiguous.
    """
    pairs = []
    for k, v in abbr_map_raw.items():
        k, v = str(k).strip(), str(v).strip()
        if not k or not v:
            continue
        if _is_acronym(k) and not _is_acronym(v):
            abbr, full = k, v
        elif _is_acronym(v) and not _is_acronym(k):
            abbr, full = v, k
        else:
            # fallbacks: prefer the token without spaces as abbr; else shorter as abbr
            if " " in k and " " not in v:
                abbr, full = v, k
            elif " " not in k and " " in v:
                abbr, full = k, v
            else:
                abbr, full = (k, v) if len(k) <= len(v) else (v, k)
        pairs.append((abbr, full))
    # deduplicate by abbr (last wins to match your original override behaviour)
    dedup = {}
    for abbr, full in pairs:
        dedup[abbr] = full
    return [(a, dedup[a]) for a in dedup]

def clean_nested_expansion(text: str):
    # Remove nested duplicates like: Full (Full (ABBR)) -> Full (ABBR)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

def handle_special_case_usa(text: str):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)
    return text

# ---------------- classification (unchanged) ----------------
def classify_topic(doc_path: str) -> str:
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    paragraph = paragraphs[1] if len(paragraphs) > 1 else ""

    categories = list(CATEGORY_TO_JSON.keys())
    system_prompt = (
        "You are an expert in classifying documents. Based only on the given title and paragraph, "
        f"classify the content into one of the following categories:\n{', '.join(categories)}.\n"
        "Respond ONLY with the exact name of the category."
    )
    user_prompt = f"Title: {title}\nParagraph: {paragraph}"

    try:
        resp = client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=100,
            temperature=0,
            system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        cat = resp.content[0].text.strip()
    except Exception:
        cat = ""
    return cat

# ---------------- quote-aware helpers (unchanged) ----------------
QUOTE_SPLIT_RE = re.compile(
    r'(".*?"|“.*?”|(?<!\w)\'.*?\'(?!\w)|(?<!\w)‘.*?’(?!\w))',
    re.DOTALL
)

def split_preserve_quotes(text: str):
    parts = []
    tokens = QUOTE_SPLIT_RE.split(text or "")
    for tok in tokens:
        if not tok:
            continue
        if QUOTE_SPLIT_RE.fullmatch(tok):
            parts.append((True, tok))
        else:
            parts.append((False, tok))
    return parts

def apply_outside_quotes(text: str, fn):
    out = []
    for is_quoted, seg in split_preserve_quotes(text):
        out.append(seg if is_quoted else fn(seg))
    return "".join(out)

# ---------------- main expand (kept logic, safer maps) ----------------
def expand_abbreviations(doc_path: str, out_path: str | None = None, base_json_path: str = "abbreviations.json") -> str:
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"File not found: {doc_path}")

    if out_path is None:
        out_path = "expanded_result.docx"

    # Load base map (accept both spellings; prefer explicit path arg if provided)
    base_map = _load_any_abbrev_map(base_json_path, "abbreviations.json", "abbrevations.json", "ABBREVIATIONS.json")

    # Classify + load category map if present
    category = classify_topic(doc_path)
    category_map = {}
    if category:
        cat_file = CATEGORY_TO_JSON.get(category)
        if cat_file:
            category_map = _load_json_first_found(cat_file)

    # Merge maps (category overrides base)
    raw = {}
    raw.update(base_map)
    raw.update(category_map)

    # Convert to canonical (abbr, full) pairs
    pairs = _to_pairs(raw)

    # Counts per ABBR (to detect first vs subsequent mentions)
    counts = defaultdict(int)

    input_doc = Document(doc_path)
    result_doc = Document()

    # Your original synonym groups
    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China", "People's Republic of China", "PRC"],
    ]

    def is_heading(paragraph) -> bool:
        try:
            return (paragraph.style.name or "").startswith("Heading")
        except Exception:
            return False

    def is_in_table(paragraph) -> bool:
        try:
            return paragraph._element.getparent().tag.endswith("tbl")
        except Exception:
            return False

    # headers/footers (skip)
    header_texts, footer_texts = [], []
    try:
        for section in input_doc.sections:
            header_texts.extend(p.text.strip() for p in section.header.paragraphs)
            footer_texts.extend(p.text.strip() for p in section.footer.paragraphs)
    except Exception:
        pass

    def is_in_header_or_footer(text: str) -> bool:
        s = (text or "").strip()
        return s in header_texts or s in footer_texts

    # -------- Pre-pass: synonym resolution outside quotes only (unchanged) --------
    raw_chunks = [
        p.text for p in input_doc.paragraphs
        if not is_heading(p) and not is_in_table(p) and not is_in_header_or_footer(p.text)
    ]
    joined = "\n".join(raw_chunks)

    def _syn_fn(s: str) -> str:
        # Keep earliest form among variants
        for group in synonym_groups:
            positions = {}
            for term in group:
                m = re.search(rf"(?<!\w){re.escape(term)}(?!\w)", s, flags=re.IGNORECASE)
                if m:
                    positions[term] = m.start()
            if positions:
                chosen = min(positions, key=positions.get)
                for alt in group:
                    if alt.lower() != chosen.lower():
                        s = re.sub(rf"(?<!\w){re.escape(alt)}(?!\w)", chosen, s, flags=re.IGNORECASE)
        return s

    full_text = apply_outside_quotes(joined, _syn_fn)

    non_skipped = full_text.split("\n")
    idx = 0
    for para in input_doc.paragraphs:
        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(para.text):
            para.text = para.text
        else:
            para.text = non_skipped[idx]
            idx += 1
    # ------------------------------------------------------------------------------

    # Names: collapse to last name after first full-name mention (kept)
    name_counts = defaultdict(int)
    name_pattern = re.compile(r'\b(?:Mr|Mrs|Ms|Dr|Prof)?\.?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b')

    def replace_names(text: str) -> str:
        matches = list(name_pattern.finditer(text or ""))
        for match in matches:
            full_name = match.group(1).strip()
            name_counts[full_name] += 1
            if name_counts[full_name] > 1:
                last_name = full_name.split()[-1]
                text = text.replace(match.group(0), last_name, 1)
        return text

    # ===== Main processing pass (only on unquoted segments) =====
    for para in input_doc.paragraphs:
        line = (para.text or "").strip()

        # keep your full-line quote skip
        if (
            (line.startswith('"') and line.endswith('"')) or
            (line.startswith("'\"\"") and line.endswith("'")) or
            (line.startswith('"\'\'') and line.endswith('"'))
        ):
            result_doc.add_paragraph(line)
            continue

        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(line):
            result_doc.add_paragraph(line)
            continue

        segments = split_preserve_quotes(line)
        out_parts = []

        for is_quoted, seg in segments:
            if is_quoted:
                out_parts.append(seg)
                continue

            seg = replace_names(seg)
            seg = clean_nested_expansion(seg)

            # --- First mention => "Full (ABBR)", then collapse Full→ABBR subsequently ---
            for abbr, full in pairs:
                # Lookaround-anchored tokens (works with dotted acronyms like "U.S.")
                abbr_pat = rf'(?<!\w){re.escape(abbr)}(?!\w)'
                full_pat = rf'(?<!\w){re.escape(full)}(?!\w)'

                # already exactly "Full (ABBR)"? (inside parentheses, anchors not needed)
                if re.search(rf'(?i){re.escape(full)}\s*\(\s*{re.escape(abbr)}\s*\)', seg):
                    counts[abbr] += 1
                    continue

                # any occurrence to act on?
                any_hit = bool(re.search(abbr_pat, seg, re.IGNORECASE) or re.search(full_pat, seg, re.IGNORECASE))
                if not any_hit:
                    continue

                counts[abbr] += 1
                if counts[abbr] == 1:
                    # first mention -> replace whichever appears first (abbr or full) with "Full (ABBR)"
                    m_abbr = re.search(abbr_pat, seg, re.IGNORECASE)
                    m_full = re.search(full_pat, seg, re.IGNORECASE)
                    if m_abbr and m_full:
                        if m_abbr.start() < m_full.start():
                            seg = re.sub(abbr_pat, f"{full} ({abbr})", seg, count=1, flags=re.IGNORECASE)
                        else:
                            seg = re.sub(full_pat, f"{full} ({abbr})", seg, count=1, flags=re.IGNORECASE)
                    elif m_abbr:
                        seg = re.sub(abbr_pat, f"{full} ({abbr})", seg, count=1, flags=re.IGNORECASE)
                    else:
                        seg = re.sub(full_pat, f"{full} ({abbr})", seg, count=1, flags=re.IGNORECASE)
                else:
                    # subsequent -> collapse remaining "Full" to ABBR (but don't touch "Full (ABBR)")
                    seg = re.sub(
                        rf'(?i){re.escape(full)}(?!\s*\(\s*{re.escape(abbr)}\s*\))',
                        abbr, seg
                    )

            seg = handle_special_case_usa(seg)
            out_parts.append(seg)

        final_line = "".join(out_parts)
        result_doc.add_paragraph(final_line)

    result_doc.save(out_path)
    return out_path

# Optional CLI
if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Expand abbreviations in a DOCX (first mention → Full (ABBR)).")
    ap.add_argument("docx", help="Input .docx path")
    ap.add_argument("--out", help="Output .docx path (optional)")
    ap.add_argument("--base-json", default="abbreviations.json", help="Base abbreviations JSON filename")
    args = ap.parse_args()
    written = expand_abbreviations(args.docx, out_path=args.out, base_json_path=args.base_json)
    print("Written:", written)
