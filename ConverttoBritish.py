#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
numbersandcurrencies_uk.py

ISAS UK style normaliser for .docx

What it does
- Convert to English (UK) spellings via breame (e.g., organize→organise, color→colour).
- Preserve official names that must use 'Organization' (e.g., World Health Organization, World Trade Organization).
- Enforce consistent foreign spellings (e.g., Ulema, Shia) across the document.
- Detect foreign / other-language words or phrases with an LLM (Claude 3.5 Haiku),
  italicise them and append a short English gloss in brackets: *term* (Gloss).

Scope
- Processes only BODY PARAGRAPHS.
- Does NOT process tables, headers, footers, footnotes/endnotes.

Setup
- pip install python-docx gradio anthropic python-dotenv breame
- create a file named: env (1)   containing:  ANTHROPIC_API_KEY=your_key_here
"""

import os
import re
import io
import json
import time
import tempfile
from typing import Dict, List, Tuple

from dotenv import load_dotenv
from docx import Document
import gradio as gr

# breame spelling helpers
from breame.spelling import (
    get_british_spelling,
    american_spelling_exists,
    british_spelling_exists
)

# Load API key from local env (1)
load_dotenv("env (1)")
ANTHROPIC_MODEL = "claude-3-5-haiku-latest"

# --------------------------------------------------
# Official names that MUST keep 'Organization'
# (extend as needed)
# --------------------------------------------------
ORG_EXCEPTIONS = {
    "World Health Organization",
    "World Trade Organization",
    "International Civil Aviation Organization",
}
ORG_EXCEPTIONS_LOWER = {s.lower(): s for s in ORG_EXCEPTIONS}

# --------------------------------------------------
# Consistent foreign spellings (left = preferred)
# (extend per project style sheet)
# --------------------------------------------------
CONSISTENT_CHOICES = {
    "Ulema": {"Uluma", "Ulama", "Ulamaa"},
    "Shia": {"Shiah", "Shi'a", "Shi‘a", "Shiite"},
}

# --------------------------------------------------
# UK spelling core via breame (tokens only)
# --------------------------------------------------
WORD_RE = re.compile(r"[A-Za-z][A-Za-z'\-]*")

def _is_exception_window(text: str, start: int, end: int) -> bool:
    window = text[max(0, start-60):min(len(text), end+60)].lower()
    for exc in ORG_EXCEPTIONS_LOWER:
        if exc in window:
            return True
    return False

def _to_uk_spelling_token(token: str) -> str:
    base = token
    caps = base.isupper()
    title = base.istitle()

    if british_spelling_exists(base.lower()):
        return base

    if american_spelling_exists(base.lower()):
        uk = get_british_spelling(base.lower()) or base
        if caps:
            uk = uk.upper()
        elif title:
            uk = uk.capitalize()
        return uk

    return base

def britishise_text_with_exceptions(text: str) -> str:
    out = []
    idx = 0
    for m in WORD_RE.finditer(text):
        start, end = m.span()
        out.append(text[idx:start])
        token = m.group(0)
        if _is_exception_window(text, start, end):
            out.append(token)
        else:
            out.append(_to_uk_spelling_token(token))
        idx = end
    out.append(text[idx:])
    result = "".join(out)

    # Restore canonical casing of exception phrases if touched
    low = result.lower()
    for exc_low, exc_canon in ORG_EXCEPTIONS_LOWER.items():
        pos = 0
        while True:
            i = low.find(exc_low, pos)
            if i == -1:
                break
            result = result[:i] + exc_canon + result[i+len(exc_canon):]
            low = result.lower()
            pos = i + len(exc_canon)

    return result

def enforce_consistent_foreign_spellings(text: str) -> str:
    for canonical, variants in CONSISTENT_CHOICES.items():
        for var in sorted(variants, key=len, reverse=True):
            def repl(m):
                s = m.group(0)
                if s.isupper():
                    return canonical.upper()
                if s.istitle():
                    return canonical
                if s.islower():
                    return canonical.lower()
                return canonical
            text = re.sub(rf"\b{re.escape(var)}\b", repl, text, flags=re.IGNORECASE)
    return text

def apply_isas_uk_rules_to_paragraph_text(text: str) -> str:
    text = britishise_text_with_exceptions(text)
    text = enforce_consistent_foreign_spellings(text)
    return text

# --------------------------------------------------
# LLM helpers (batched) for foreign / other-language terms
# --------------------------------------------------
def _looks_foreign(text: str) -> bool:
    if re.search(r"[^\x00-\x7F]", text):  # non-ASCII
        return True
    cues = ["Abhiyan", "Yatra", "Sangh", "Sabha", "ayatollah", "madrasa", "shura", "qazi", "ulema"]
    for w in cues:
        if re.search(rf"\b{re.escape(w)}\b", text, flags=re.IGNORECASE):
            return True
    return False

def _batch_paragraphs_for_gloss(paragraphs: List[str], max_chars: int = 8000) -> List[List[Tuple[int, str]]]:
    chunks: List[List[Tuple[int, str]]] = []
    cur: List[Tuple[int, str]] = []
    cur_len = 0
    for idx, txt in enumerate(paragraphs):
        if not txt or not _looks_foreign(txt):
            continue
        tlen = len(txt) + 50
        if cur and cur_len + tlen > max_chars:
            chunks.append(cur)
            cur = []
            cur_len = 0
        cur.append((idx, txt))
        cur_len += tlen
    if cur:
        chunks.append(cur)
    return chunks

def _llm_gloss_batch(candidates: List[Tuple[int, str]]) -> Dict[int, List[Dict[str, str]]]:
    import anthropic
    if not candidates:
        return {}
    payload = [{"id": idx, "text": txt} for idx, txt in candidates]
    prompt = (
        "You are an editorial assistant for UK academic style.\n"
        "For each paragraph item, identify up to 6 non-English (foreign/other-language) words or short phrases "
        "that should be italicised, and provide a concise English gloss (≤ 5 words). "
        "Return STRICT JSON only as:\n"
        "{ \"results\": [ {\"id\": int, \"items\": [ {\"phrase\": str, \"gloss\": str}, ...]}, ...] }\n"
        "Only include clearly foreign terms; no extra fields or commentary.\n"
        f"Items:\n{json.dumps(payload, ensure_ascii=False)}"
    )
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=1500,
        temperature=0,
        messages=[{"role":"user","content":prompt}],
    )
    text = "".join(getattr(b, "text", "") for b in msg.content)
    try:
        data = json.loads(text)
        results = data.get("results", [])
        out: Dict[int, List[Dict[str, str]]] = {}
        for entry in results:
            pid = entry.get("id")
            items = entry.get("items", [])
            if isinstance(pid, int) and isinstance(items, list):
                clean = []
                for it in items:
                    phrase = (it.get("phrase") or "").strip()
                    gloss = (it.get("gloss") or "").strip()
                    if phrase and gloss:
                        clean.append({"phrase": phrase, "gloss": gloss})
                out[pid] = clean[:6]
        return out
    except Exception:
        return {}

def italicise_and_gloss(paragraph, items: List[Dict[str, str]]):
    if not items:
        return
    items_sorted = sorted(items, key=lambda d: len(d["phrase"]), reverse=True)
    for it in items_sorted:
        phrase = it["phrase"]
        gloss = it["gloss"]
        if not phrase or not gloss:
            continue
        # simple per-run replace if phrase fully inside this run
        for run in paragraph.runs:
            if not run.text:
                continue
            i = run.text.find(phrase)
            if i == -1:
                continue
            before = run.text[:i]
            after = run.text[i+len(phrase):]
            run.text = before
            phr_run = paragraph.add_run(phrase)
            phr_run.italic = True
            paragraph.add_run(f" ({gloss})")
            paragraph.add_run(after)

# --------------------------------------------------
# DOCX processing (BODY ONLY)
# --------------------------------------------------
def process_document(input_bytes: bytes) -> bytes:
    # save input to temp path for python-docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
        tmp_in.write(input_bytes)
        in_path = tmp_in.name

    doc = Document(in_path)

    # 1) Apply UK spelling & consistency (local, fast)
    for p in doc.paragraphs:
        original = p.text
        updated = apply_isas_uk_rules_to_paragraph_text(original)
        if updated != original:
            if p.runs:
                p.runs[0].text = updated
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(updated)

    # 2) LLM: detect foreign terms (batched), then italicise + gloss
    paragraphs_text = [p.text for p in doc.paragraphs]
    chunks = _batch_paragraphs_for_gloss(paragraphs_text, max_chars=8000)
    # optional cache if repeated phrases appear
    phrase_cache: Dict[str, str] = {}

    for chunk in chunks:
        result_map = _llm_gloss_batch(chunk)  # {idx: [{"phrase","gloss"}, ...]}
        for idx, items in result_map.items():
            # cache & dedupe
            dedup: List[Dict[str, str]] = []
            for it in items:
                ph = it["phrase"]
                gl = it["gloss"]
                if ph in phrase_cache:
                    gl = phrase_cache[ph]
                else:
                    phrase_cache[ph] = gl
                dedup.append({"phrase": ph, "gloss": gl})
            italicise_and_gloss(doc.paragraphs[idx], dedup)

    # 3) Write output
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

# --------------------------------------------------
# Gradio UI (no checkboxes; body-only)
# --------------------------------------------------
def convert_file(docx_path):
    if not docx_path:
        raise gr.Error("Please upload a .docx file.")
    with open(str(docx_path), "rb") as f:
        data = f.read()
    output_bytes = process_document(data)
    out_path = os.path.join(tempfile.gettempdir(), f"isas_uk_style_{int(time.time())}.docx")
    with open(out_path, "wb") as f:
        f.write(output_bytes)
    return out_path

with gr.Blocks(title="ISAS UK Style (.docx)") as demo:
    gr.Markdown(
        "## ISAS UK Style Normaliser (Body paragraphs only)\n"
        "- British English via breame (with exceptions for official names).\n"
        "- Consistent foreign spellings (e.g., Ulema, Shia).\n"
        "- Foreign / other-language terms italicised with a short English gloss (LLM).\n\n"
        "**Note:** Tables, headers, footers, and notes are not processed."
    )
    inp = gr.File(label="Upload .docx", file_types=[".docx"], type="filepath")
    out = gr.File(label="Download processed .docx")
    go = gr.Button("Apply UK style")
    go.click(convert_file, inputs=[inp], outputs=out)

if __name__ == "__main__":
    demo.launch()
