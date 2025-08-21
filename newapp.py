# app.py — One-tool pipeline:
#   1) Style pass (deterministic rules + Claude with system/user split; strict "Corrected:" extraction)
#   2) Footnotes pass (footnote_corrector.py using fcorrectprompt.py)
#   3) Abbreviations pass LAST (test8.expand_abbreviations)
#   4) Diff-highlighting — compares original vs final and marks changes in yellow (body + footnotes/endnotes)
#
# UI: upload -> download .docx; share=True. Default model: claude-3-7-sonnet-latest.

from typing import Optional, Tuple, List, Any
import os, re, time, importlib.util, unicodedata
import gradio as gr
from difflib import SequenceMatcher
from itertools import zip_longest
from shutil import copyfile

# ---------------- ENV ----------------
try:
    from dotenv import load_dotenv
except Exception:
    load_dotenv = None

if load_dotenv:
    try:
        load_dotenv()
        if os.path.exists("env (1)"):
            load_dotenv("env (1)", override=True)
    except Exception:
        pass

# ---------------- Anthropic ----------------
try:
    from anthropic import Anthropic
except Exception:
    Anthropic = None  # type: ignore

# ---------------- DOCX ----------------
try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.text.paragraph import Paragraph  # for footnote/endnote paragraphs
except Exception as e:
    raise RuntimeError("python-docx missing. Install: pip install python-docx") from e

# ---------------- Optional UK spelling (Breame) ----------------
try:
    from breame.spelling import get_british_spelling
except Exception:
    get_british_spelling = None

# ---------------- Your big style prompt ----------------
from updated_prompt import MAJOR_PROMPT

# ======================================================================================
# CONFIG
# ======================================================================================
DEFAULT_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-3-7-sonnet-latest")
FALLBACK_MODEL = os.getenv("ANTHROPIC_FALLBACK_MODEL", "claude-3-5-sonnet-latest")
MAX_TOKENS = int(os.getenv("ANTHROPIC_MAX_TOKENS", "2048"))

def _anthropic_client() -> Optional[Any]:
    if Anthropic is None:
        return None
    api_key = (
        os.getenv("ANTHROPIC_API_KEY")
        or os.getenv("ANTHROPIC_KEY")
        or os.getenv("ANTHROPIC_APIKEY")
    )
    if not api_key:
        return None
    try:
        return Anthropic(api_key=api_key)
    except Exception:
        return None

# ======================================================================================
# MASKING / HELPERS
# ======================================================================================
def _merge_spans(spans: List[Tuple[int,int]]) -> List[Tuple[int,int]]:
    if not spans: return []
    spans = sorted(spans)
    merged = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return [(s, e) for s, e in merged]

def _idx_to_letters(i: int) -> str:
    if i < 0: i = 0
    s = ""
    i += 1
    while i > 0:
        i, r = divmod(i - 1, 26)
        s = chr(ord('a') + r) + s
    return s

def _make_token(tag: str, i: int) -> str:
    return f"§{tag}§{_idx_to_letters(i)}§"

def _mask_ranges(text: str, spans: List[Tuple[int,int]], tag: str) -> Tuple[str, List[str]]:
    spans = _merge_spans(spans)
    buckets, out, last = [], [], 0
    for _, (s, e) in enumerate(spans):
        out.append(text[last:s])
        out.append(_make_token(tag, len(buckets)))
        buckets.append(text[s:e])
        last = e
    out.append(text[last:])
    return ''.join(out), buckets

def _unmask(text: str, tag: str, buckets: List[str]) -> str:
    for i, v in enumerate(buckets):
        text = text.replace(_make_token(tag, i), v)
    return text

def _find_code_fences(text: str):
    return [(m.start(), m.end()) for m in re.finditer(r'```.*?\n.*?```', text, flags=re.DOTALL)]

def _find_quoted_segments(text: str):
    spans = []
    for m in re.finditer(r'"[^"\n]*"', text): spans.append((m.start(), m.end()))
    for m in re.finditer(r'“[^”\n]*”', text): spans.append((m.start(), m.end()))
    for m in re.finditer(r"(?<!\w)'[^'\n]*'(?!\w)", text): spans.append((m.start(), m.end()))
    for m in re.finditer(r"(?<!\w)‘[^’\n]*’(?!\w)", text): spans.append((m.start(), m.end()))
    return _merge_spans(spans)

def _is_table_or_figure_line(line: str) -> bool:
    if line.count('|') >= 2: return True
    if line.count('\t') >= 2: return True
    if re.search(r'^\s*(Table|Figure|Fig\.|Chart|Exhibit)\b', line, re.IGNORECASE): return True
    if re.match(r'^\s*\|? *-+ *(?:\| *-+ *)+\|?\s*$', line): return True
    return False

def _find_html_blocks(text: str, tags):
    spans = []
    for tag in tags:
        pattern = re.compile(rf'<{tag}\b[^>]*>.*?</{tag}>', re.IGNORECASE | re.DOTALL)
        spans += [(m.start(), m.end()) for m in pattern.finditer(text)]
    return spans

def _find_html_single_tags(text: str, tag: str):
    pat = re.compile(rf'<{tag}\b[^>]*>', re.IGNORECASE | re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_heading_spans(text: str):
    spans = []
    for m in re.finditer(r'(?m)^(#{1,6})\s.*$', text): spans.append((m.start(), m.end()))
    for m in re.finditer(r'(?ms)^(?P<h>.+?)\n(=+|-{3,})\s*$', text): spans.append((m.start('h'), m.end()))
    return spans

def _find_md_image_spans(text: str):
    pat = re.compile(r'!\[[^\]]*\]\([^)]+\)', re.DOTALL)
    return [(m.start(), m.end()) for m in pat.finditer(text)]

def _find_md_footnote_block_spans(text: str):
    spans = []
    for m in re.finditer(r'(?m)^\[\^[^\]]+\]:[^\n]*\n(?:[ \t].*\n)*', text):
        spans.append((m.start(), m.end()))
    return spans

def _find_docx_placeholder_spans(text: str):
    pats = [r'§FOOTNOTE:[^§]+§', r'FN_PLACEHOLDER[^§\n]*']
    spans = []
    for p in pats:
        for m in re.finditer(p, text):
            spans.append((m.start(), m.end()))
    return spans

def _find_asterisk_marked_spans(text: str):
    spans = []
    for m in re.finditer(r'\*\*.*?\*\*', text, flags=re.DOTALL): spans.append((m.start(), m.end()))
    for m in re.finditer(r'\*(?!\*)(.*?)\*', text, flags=re.DOTALL): spans.append((m.start(), m.end()))
    return _merge_spans(spans)

def _find_par_markers(text: str):
    return [(m.start(), m.end()) for m in re.finditer(re.escape("<<par>>"), text)]

def _collect_protected_spans(text: str):
    spans = []
    spans += _find_code_fences(text)
    spans += _find_quoted_segments(text)
    spans += _find_html_blocks(text, ['table','thead','tbody','tfoot','tr','header','footer','figure'])
    spans += _find_html_single_tags(text, 'img')
    spans += _find_md_heading_spans(text)
    spans += _find_md_image_spans(text)
    spans += _find_md_footnote_block_spans(text)
    spans += _find_docx_placeholder_spans(text)
    return _merge_spans(spans)

def _xml_safe_text(s: str) -> str:
    if s is None: return ""
    s = unicodedata.normalize("NFC", s)
    s = s.replace("\x0b", "\n").replace("\x0c", "\n")
    s = re.sub(r"[\x00-\x08\x0E-\x1F]", "", s)
    s = re.sub(r"[\uFFFE\uFFFF]", "", s)
    return s

# ======================================================================================
# NUMBERS / CURRENCY / UK (style rules)
# ======================================================================================
def _apply_currency_rules(text: str) -> str:
    CODES = {'GBP':'£','EUR':'€','JPY':'¥','SGD':'S$','USD':'US$'}
    for code, sym in CODES.items():
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)\b', rf'{sym}\1', text, flags=re.IGNORECASE)
        text = re.sub(rf'\b{code}\s*([0-9][\d,]*(?:\.\d+)?)(\s+)(million|billion)\b',
                      rf'{sym}\1 \3', text, flags=re.IGNORECASE)
    text = re.sub(r'(?:(?<=€)|(?<=£)|(?<=¥)|(?<=₹)|(?<=S\$)|(?<=US\$))\s+(?=\d)', '', text)
    text = re.sub(r'((?:€|£|¥|₹|S\$|US\$)\d[\d,]*(?:\.\d+)?)(\s*)(million|billion)\b', r'\1 \3', text, flags=re.IGNORECASE)
    return text

def _mask_currencies(text: str):
    pat = re.compile(r'(?:US\$|S\$|€|£|¥|₹)\s?\d[\d,]*(?:\.\d+)?(?:\s+(?:million|billion))?', re.IGNORECASE)
    spans = [(m.start(), m.end()) for m in pat.finditer(text)]
    return _mask_ranges(text, spans, 'CUR')

_WORDS_0_19 = ["zero","one","two","three","four","five","six","seven","eight","nine",
               "ten","eleven","twelve","thirteen","fourteen","fifteen","sixteen",
               "seventeen","eighteen","nineteen"]
_TENS = ["","","twenty","thirty","forty","fifty","sixty","seventy","eighty","ninety"]

def _int_to_words(n: int) -> str:
    if n < 20: return _WORDS_0_19[n]
    if n < 100:
        tens, rem = divmod(n, 10)
        return _TENS[tens] + ("" if rem == 0 else "-" + _WORDS_0_19[rem])
    return str(n)

def _decimal_to_words_token(tok: str) -> str:
    s = tok.strip()
    neg = s.startswith('-') or s.startswith('−')
    if neg: s = s[1:]
    if s.startswith('.'):
        int_part, frac_part = 0, s[1:]
    else:
        int_part, frac_part = s.split('.', 1)
    int_words = _int_to_words(int(int_part)) if int_part else "zero"
    frac_words = " ".join(_WORDS_0_19[int(d)] for d in frac_part if d.isdigit())
    return f"{'minus ' if neg else ''}{int_words} point {frac_words}"

_STAY_IZE = {"capsize","capsizes","capsized","capsizing","seize","seizes","seized","seizing",
             "prize","prizes","prized","prizing","size","sizes","sized","sizing"}
_IZE_SUFFIX_RE = re.compile(r'(?:ize|izes|ized|izing|ization|izations)\b')
_WORD_LOWER_RE = re.compile(r'\b([a-z]+)\b')

def _detect_preserve_ize_pref(raw_text: str) -> bool:
    if not get_british_spelling: return False
    ize = len(re.findall(r'\b[a-z]+(?:ize|izes|ized|izing|ization|izations)\b', raw_text))
    ise = len(re.findall(r'\b[a-z]+(?:ise|ises|ised|ising|isation|isations)\b', raw_text))
    return ize >= 5 and ize >= max(8, int(1.8 * max(1, ise)))

def _apply_breame_uk_to_line(line: str, preserve_ize: bool) -> str:
    if get_british_spelling is None: return line
    def repl(m):
        w = m.group(1)
        if w in _STAY_IZE: return w
        uk = get_british_spelling(w)
        if not uk or uk == w: return w
        if preserve_ize and _IZE_SUFFIX_RE.search(w): return w
        return uk
    return _WORD_LOWER_RE.sub(repl, line)

# ======================================================================================
# STYLE PIPELINE (deterministic rules + Claude with MAJOR_PROMPT, no prompt bleed)
# ======================================================================================
def _spell_sentence_start_integers(line: str) -> str:
    TOKEN_RE = r'(?:§(?:PROT|CUR|AST|PAR)§[a-z]+§)'
    SENT_START_INT = re.compile(rf'(^|(?<=[.!?])(?:\s|{TOKEN_RE})+)(\d{{1,3}}(?:,\d{{3}})*)\b(?!\.\d)')
    def _repl(m):
        prefix = m.group(1)
        num_str = m.group(2)
        n = int(num_str.replace(',', ''))
        words = _int_to_words(n) if 0 <= n <= 9 else str(n)
        return prefix + words[:1].upper() + words[1:]
    return SENT_START_INT.sub(_repl, line)

def apply_all_rules(text: str) -> str:
    text = _apply_currency_rules(text)
    preserve_ize = _detect_preserve_ize_pref(text)

    prot_spans = _collect_protected_spans(text)
    text, prot_bucket = _mask_ranges(text, prot_spans, 'PROT')

    ast_spans = _find_asterisk_marked_spans(text)
    text, ast_bucket = _mask_ranges(text, ast_spans, 'AST')

    par_spans = _find_par_markers(text)
    text, par_bucket = _mask_ranges(text, par_spans, 'PAR')

    text, cur_bucket = _mask_currencies(text)

    lines = text.splitlines(keepends=False)
    for i, line in enumerate(lines):
        if _is_table_or_figure_line(line):
            lines[i] = line
            continue

        line = _apply_breame_uk_to_line(line, preserve_ize)

        def _spell_small(n: int) -> str:
            return _WORDS_0_19[n] if 0 <= n <= 9 else str(n)

        # Ranges 3–4% -> three to four per cent; 3–4 -> three to four
        line = re.sub(
            r'(?<![\d.,])(\d{1,3})\s*[-–]\s*(\d{1,3})\s*(?:%|\uFF05)(?!\w)',
            lambda m: f"{_spell_small(int(m.group(1)))} to {_spell_small(int(m.group(2)))} per cent",
            line
        )
        line = re.sub(
            r'(?<![\d.,])(\d{1,3})\s*[-–]\s*(\d{1,3})(?!\s*(?:%|\uFF05))',
            lambda m: f"{_spell_small(int(m.group(1)))} to {_spell_small(int(m.group(2)))}",
            line
        )

        # Integers with % -> "N per cent" (decimals with % stay as-is)
        line = re.sub(r'(?<![\d.,])(\d{1,3}(?:,\d{3})*)\s*(?:%|\uFF05)([A-Za-z])', r'\1 per cent \2', line)
        line = re.sub(r'(?<![\d.,])(\d{1,3}(?:,\d{3})*)\s*(?:%|\uFF05)(?!\w)', r'\1 per cent', line)
        # x percent -> x per cent (decimals allowed)
        line = re.sub(
            r'(?<![\w§])(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\.\d+)\s*percent\b(?=[\s\.,;:!\?\)\]»›”’]|$)',
            r'\1 per cent', line, flags=re.IGNORECASE
        )
        # word-number percent -> N per cent
        line = re.sub(
            r'\b('
            r'zero|one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|'
            r'fifteen|sixteen|seventeen|eighteen|nineteen|twenty|thirty|forty|fifty|sixty|seventy|'
            r'eighty|ninety(?:-(?:one|two|three|four|five|six|seven|eight|nine))?'
            r')\s*percent\b(?=[\s\.,;:!\?\)\]]|$)',
            lambda m: f"{_wordnum_to_int(m.group(1))} per cent", line, flags=re.IGNORECASE
        )

        # Temperature -> "Degree" text
        def temp_repl(m):
            val, unit = m.group(1), m.group(2).upper()
            return f"{val} Degree {'Celsius' if unit=='C' else 'Fahrenheit'}"
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*°\s*([CF])\b', temp_repl, line)
        line = re.sub(r'(?<![\w-])([+-−]?\d+(?:\.\d+)?)\s*([CF])\b(?![-\d])', temp_repl, line)

        # Decimals -> words, unless followed by %/percent/temperature
        def dec_repl(m):
            tok = m.group(1)
            after = line[m.end():]
            if re.match(r'^\s*(?:per\s*cent\b|percent\b|%|\uFF05)', after, flags=re.IGNORECASE): return tok
            if re.match(r'^\s*(?:°\s*[CF]|[CF]\b|degree\s+(?:celsius|fahrenheit))', after, flags=re.IGNORECASE): return tok
            return _decimal_to_words_token(tok)
        line = re.sub(r'(?<![\w./:\-§])((?:\d+\.\d+)|(?:\.\d+))(?!\.\d)(?![\w./:\-§])', dec_repl, line)

        # Single digits -> words (standalone)
        line = re.sub(r'(?<![\w./:\-§])([0-9])(?![\w./:\-§])', lambda m: _WORDS_0_19[int(m.group(1))], line)

        # Sentence-start smaller integers
        line = _spell_sentence_start_integers(line)
        lines[i] = line

    text = '\n'.join(lines)
    text = _unmask(text, 'CUR', cur_bucket)
    text = _unmask(text, 'PAR', par_bucket)
    text = _unmask(text, 'AST', ast_bucket)
    text = _unmask(text, 'PROT', prot_bucket)
    return text

def _wordnum_to_int(word: str) -> int:
    m = {"zero":0,"one":1,"two":2,"three":3,"four":4,"five":5,"six":6,"seven":7,"eight":8,"nine":9,
         "ten":10,"eleven":11,"twelve":12,"thirteen":13,"fourteen":14,"fifteen":15,"sixteen":16,
         "seventeen":17,"eighteen":18,"nineteen":19,
         "twenty":20,"thirty":30,"forty":40,"fifty":50,"sixty":60,"seventy":70,"eighty":80,"ninety":90}
    w = word.lower()
    if '-' in w:
        a, b = w.split('-', 1)
        return m[a] + m[b]
    return m[w]

# --------- Prompt split + Anthropic call (no prompt bleed) ----------
def _split_prompt_for_system_user(major_prompt: str, masked_text: str) -> tuple[str, str]:
    head, sep, _ = major_prompt.partition("Input text:")
    system_prompt = head.strip() if sep else major_prompt.replace("{text}", "").strip()
    user_msg = masked_text
    return system_prompt, user_msg

def _anthropic_complete(system_prompt: str, user_text: str, model: str, max_tokens: int) -> Tuple[Optional[str], Optional[str]]:
    client = _anthropic_client()
    if client is None:
        return None, "Anthropic client not available (missing SDK or API key)."
    last_err = None
    for mdl in (model, FALLBACK_MODEL if FALLBACK_MODEL != model else None):
        if not mdl: continue
        try:
            msg = client.messages.create(
                model=mdl,
                system=system_prompt,
                messages=[{"role": "user", "content": user_text}],
                temperature=0,
                max_tokens=max_tokens,
                stop_sequences=["ROLE","TASK","CHAIN-OF-THOUGHT","DO NOT","OUTPUT FORMAT","Input text:"],
            )
            out = "".join(block.text for block in msg.content if getattr(block, "type", "text") == "text")
            return out, None
        except Exception as e:
            last_err = str(e)
    return None, f"Anthropic error: {last_err}"

def _extract_corrected_text(raw: str) -> Optional[str]:
    if raw is None: return None
    m = re.search(r'^\s*Corrected:\s*\n?', raw, flags=re.IGNORECASE | re.MULTILINE)
    if not m: return None
    return raw[m.end():].strip()

def apply_major_prompt_with_masks(text: str, model: str, max_tokens: int) -> Tuple[str, Optional[str]]:
    prot_spans = _collect_protected_spans(text)
    masked, prot_bucket = _mask_ranges(text, prot_spans, 'PROT')
    ast_spans = _find_asterisk_marked_spans(masked)
    masked, ast_bucket = _mask_ranges(masked, ast_spans, 'AST')
    par_spans = _find_par_markers(masked)
    masked, par_bucket = _mask_ranges(masked, par_spans, 'PAR')

    system_prompt, user_msg = _split_prompt_for_system_user(MAJOR_PROMPT, masked)
    out, err = _anthropic_complete(system_prompt, user_msg, model=DEFAULT_MODEL, max_tokens=MAX_TOKENS)
    if err or out is None:
        return text, err or "Unknown Anthropic error."

    corrected = _extract_corrected_text(out)
    if corrected is None:
        return text, "Model response missing 'Corrected:'; original text kept."

    if any(h in corrected for h in ("ROLE","TASK","CHAIN-OF-THOUGHT","DO NOT","OUTPUT FORMAT","Input text:")):
        return text, "Filtered instructional content; original text kept."

    corrected = _unmask(corrected, 'PAR', par_bucket)
    corrected = _unmask(corrected, 'AST', ast_bucket)
    corrected = _unmask(corrected, 'PROT', prot_bucket)
    return corrected, None

# ---------------- DOCX helpers for style pass ----------------
def _para_in_table(paragraph) -> bool:
    try:
        return bool(paragraph._element.xpath("ancestor::*[local-name()='tbl']"))
    except Exception:
        return False

def _is_heading(paragraph) -> bool:
    try:
        name = (paragraph.style.name or "").lower()
    except Exception:
        name = ""
    return name.startswith("heading") or name in {"title", "subtitle"}

def _is_quote_style(paragraph) -> bool:
    try:
        name = (paragraph.style.name or "").lower()
    except Exception:
        name = ""
    return "quote" in name

def _run_has_nontext_children(run) -> bool:
    r = run._r
    try:
        return bool(r.xpath(".//*[local-name()='footnoteReference' or local-name()='endnoteReference' or local-name()='fldChar' or local-name()='instrText' or local-name()='drawing']"))
    except Exception:
        return False

def _transform_paragraph_runs_grouped(paragraph, transform_fn):
    runs = paragraph.runs
    n = len(runs)
    i = 0
    while i < n:
        if _run_has_nontext_children(runs[i]):
            i += 1
            continue
        group_idx, group_txt = [], []
        while i < n and not _run_has_nontext_children(runs[i]):
            t = runs[i].text or ""
            group_idx.append(i); group_txt.append(t); i += 1
        combined = "".join(group_txt)
        if combined.strip() == "":
            continue
        new_combined = transform_fn(combined)
        if new_combined == combined:
            continue
        pos = 0
        for k, ridx in enumerate(group_idx):
            orig_len = len(group_txt[k])
            if k < len(group_idx) - 1:
                piece = new_combined[pos:pos + orig_len]
                runs[ridx].text = _xml_safe_text(piece)
                pos += orig_len
            else:
                runs[ridx].text = _xml_safe_text(new_combined[pos:])

def process_docx_style(src_path: str, dst_path: str, model: str, max_tokens: int):
    doc = Document(src_path)
    for p in doc.paragraphs:
        if _para_in_table(p) or _is_heading(p) or _is_quote_style(p):
            continue
        _transform_paragraph_runs_grouped(p, apply_all_rules)
        para_text = p.text
        if para_text.strip():
            edited, err = apply_major_prompt_with_masks(para_text, model=model, max_tokens=max_tokens)
            if not err and edited and edited != para_text:
                if any(s in edited for s in ("ROLE","TASK","CHAIN-OF-THOUGHT","DO NOT","OUTPUT FORMAT","Input text:")):
                    pass  # drop suspicious; keep original
                else:
                    for r in p.runs: r.text = ""
                    p.add_run(_xml_safe_text(edited))
    doc.save(dst_path)

# ======================================================================================
# LOAD USER MODULES — exact filenames
# ======================================================================================
def _load_module_from_path(mod_name: str, path: str):
    spec = importlib.util.spec_from_file_location(mod_name, path)
    if not spec or not spec.loader: return None
    module = importlib.util.module_from_spec(spec)
    try:
        spec.loader.exec_module(module)  # type: ignore
        return module
    except Exception:
        return None

def _load_test8_safely(mod_path: str):
    # prevent test8’s Gradio UI from launching when imported
    try:
        import gradio as _gr
        orig_launch = getattr(_gr.Interface, "launch", None)
        if callable(orig_launch):
            _gr.Interface.launch = lambda *a, **k: None
    except Exception:
        orig_launch = None
        _gr = None
    try:
        return _load_module_from_path("user_test8", mod_path)
    finally:
        try:
            if orig_launch:
                _gr.Interface.launch = orig_launch  # type: ignore
        except Exception:
            pass

def _find_user_modules():
    cwd = os.getcwd()

    # test8.py (abbrev)
    test8 = None
    for p in [os.path.join(cwd, "test8.py"), "/mnt/data/test8.py"]:
        if os.path.exists(p):
            test8 = _load_test8_safely(p)
            if test8: break

    # fcorrectprompt.py (exact)
    fprompt = None
    for p in [os.path.join(cwd, "fcorrectprompt.py"), "/mnt/data/fcorrectprompt.py"]:
        if os.path.exists(p):
            fprompt = _load_module_from_path("user_fprompt", p)
            if fprompt: break

    # footnote_corrector.py (exact)
    ftool = None
    for p in [os.path.join(cwd, "footnote_corrector.py"), "/mnt/data/footnote_corrector.py"]:
        if os.path.exists(p):
            ftool = _load_module_from_path("user_fcorrector", p)
            if ftool: break

    return test8, fprompt, ftool

def _run_footnotes_step(ftool_module, fprompt_module, src_path: str, tmp_out: str) -> str:
    if not ftool_module: return src_path
    prompt = getattr(fprompt_module, "FOOTNOTE_PROMPT", None) if fprompt_module else None
    try:
        for fn_name in ["correct_footnotes_docx", "process_docx", "run"]:
            fn = getattr(ftool_module, fn_name, None)
            if callable(fn):
                try:
                    if prompt is not None and fn.__code__.co_argcount >= 3:
                        fn(src_path, tmp_out, prompt)
                    else:
                        fn(src_path, tmp_out)
                    return tmp_out if os.path.exists(tmp_out) else src_path
                except TypeError:
                    try:
                        fn(src_path, tmp_out)
                        return tmp_out if os.path.exists(tmp_out) else src_path
                    except Exception:
                        continue
                except Exception:
                    continue
    except Exception:
        pass
    return src_path

def _run_abbrev_step_last(test8_module, src_path: str, tmp_out: str) -> str:
    """
    Run abbreviations LAST. Introspect to avoid passing tmp_out where a JSON path is expected.
    """
    if not test8_module:
        return src_path
    try:
        import inspect
        candidates = [
            "expand_abbreviations",
            "expand_abbreviations_docx",
            "process_docx",
            "process_file",
            "run",
        ]
        for name in candidates:
            fn = getattr(test8_module, name, None)
            if not callable(fn):
                continue
            sig = inspect.signature(fn)
            params = list(sig.parameters.values())
            arity = len(params)

            def _looks_like_json(p):
                n = p.name.lower()
                return "json" in n or "map" in n

            def _looks_like_out(p):
                n = p.name.lower()
                return "out" in n or "output" in n or "dst" in n or "dest" in n

            variants = []
            if name == "expand_abbreviations":
                if arity >= 2 and _looks_like_out(params[1]):
                    variants = [(fn, (src_path, tmp_out), {})]
                elif arity >= 2 and _looks_like_json(params[1]):
                    variants = [(fn, (src_path,), {})]
                else:
                    variants = [(fn, (src_path,), {})]
            else:
                if arity >= 2 and _looks_like_out(params[1]):
                    variants = [(fn, (src_path, tmp_out), {})]
                elif arity == 1:
                    variants = [(fn, (src_path,), {})]
                else:
                    variants = [(fn, (src_path,), {}), (fn, (src_path, tmp_out), {})]

            for func, args, kwargs in variants:
                try:
                    ret = func(*args, **kwargs)
                    if isinstance(ret, str) and os.path.exists(ret):
                        return ret
                    if os.path.exists(tmp_out):
                        return tmp_out
                    return src_path
                except TypeError:
                    continue
    except Exception:
        pass
    return src_path

# ======================================================================================
# TEXT I/O helpers
# ======================================================================================
def _save_text_as_docx(text: str, out_path: str):
    from docx import Document as _Doc
    doc = _Doc()
    for line in text.splitlines() or [""]:
        doc.add_paragraph(_xml_safe_text(line))
    doc.save(out_path)

def _read_text_from_path(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read()
    try: return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="ignore")

# ======================================================================================
# Diff highlighting helpers (mark changes in YELLOW in body + footnotes/endnotes)
# ======================================================================================
_TOKEN_RE = re.compile(r'\w+|\s+|[^\w\s]', re.UNICODE)

def _tokenize_for_diff(s: str) -> List[str]:
    return _TOKEN_RE.findall(s or "")

def _para_has_nontext(paragraph) -> bool:
    try:
        for r in paragraph.runs:
            if _run_has_nontext_children(r):
                return True
        return False
    except Exception:
        return True

def _paragraph_is_safe_for_rewrite(p) -> bool:
    return (not _para_in_table(p)) and (not _is_heading(p)) and (not _is_quote_style(p)) and (not _para_has_nontext(p))

def _segments_from_diff(orig_text: str, new_text: str) -> List[Tuple[str, bool]]:
    """
    Returns [(segment_text, highlight_bool)], where highlight=True for inserted/replaced text.
    """
    a = _tokenize_for_diff(orig_text)
    b = _tokenize_for_diff(new_text)
    sm = SequenceMatcher(None, a, b, autojunk=False)
    out: List[Tuple[str, bool]] = []
    buf = ""
    buf_hl: Optional[bool] = None

    def flush():
        nonlocal buf, buf_hl, out
        if buf:
            out.append((buf, bool(buf_hl)))
            buf = ""
            buf_hl = None

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for tok in b[j1:j2]:
                if buf_hl is not False:
                    flush()
                    buf_hl = False
                buf += tok
        elif tag in ("replace", "insert"):
            for tok in b[j1:j2]:
                if buf_hl is not True:
                    flush()
                    buf_hl = True
                buf += tok
        elif tag == "delete":
            # text removed from original — nothing to add to final
            continue
    flush()
    return out

def _set_paragraph_with_segments(paragraph, segments: List[Tuple[str, bool]]):
    # Clear all existing text (keep paragraph container)
    for r in paragraph.runs:
        r.text = ""
    # Add new runs with optional yellow highlight
    for text, highlight in segments:
        if not text:
            continue
        run = paragraph.add_run(text)
        if highlight:
            try:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            except Exception:
                pass

def _iter_note_paragraphs(doc, kind: str = "footnotes"):
    """
    Yield python-docx Paragraph objects from footnotes or endnotes part.
    kind: "footnotes" or "endnotes"
    """
    part = getattr(doc._part, f"{kind}_part", None)
    if not part:
        return
    ns = part.element.nsmap
    query = ".//w:footnote" if kind == "footnotes" else ".//w:endnote"
    for fn in part.element.xpath(query, namespaces=ns):
        for p in fn.xpath(".//w:p", namespaces=ns):
            yield Paragraph(p, part)

def _highlight_note_differences(doc_orig, doc_final, kind: str):
    orig_ps = list(_iter_note_paragraphs(doc_orig, kind))
    final_ps = list(_iter_note_paragraphs(doc_final, kind))
    for p_orig, p_final in zip_longest(orig_ps, final_ps, fillvalue=None):
        if p_final is None:
            continue
        # Reuse same safety checks used for body paragraphs
        if _para_has_nontext(p_final):
            continue
        orig_text = p_orig.text if p_orig is not None else ""
        final_text = p_final.text or ""
        if orig_text == final_text:
            continue
        segs = _segments_from_diff(orig_text, final_text)
        if segs and any(h for _, h in segs):
            _set_paragraph_with_segments(p_final, segs)

def highlight_doc_differences(orig_docx_path: str, final_docx_path: str):
    """
    Compare original vs final DOCX and highlight changes (inserted/replaced text) in YELLOW
    directly IN the final document.
    - Body paragraphs (safe ones)
    - Footnotes and Endnotes
    """
    try:
        doc_orig = Document(orig_docx_path)
        doc_final = Document(final_docx_path)
    except Exception:
        return

    # Body paragraphs
    for p_orig, p_final in zip_longest(doc_orig.paragraphs, doc_final.paragraphs, fillvalue=None):
        if p_final is None:
            continue
        if not _paragraph_is_safe_for_rewrite(p_final):
            continue
        orig_text = p_orig.text if p_orig is not None else ""
        final_text = p_final.text or ""
        if orig_text == final_text:
            continue
        segments = _segments_from_diff(orig_text, final_text)
        if segments and any(h for _, h in segments):
            _set_paragraph_with_segments(p_final, segments)

    # Footnotes + Endnotes
    try:
        _highlight_note_differences(doc_orig, doc_final, "footnotes")
        _highlight_note_differences(doc_orig, doc_final, "endnotes")
    except Exception:
        pass

    # Save in-place
    try:
        doc_final.save(final_docx_path)
    except Exception:
        pass

# ======================================================================================
# ONE-TOOL PIPELINE  (Style -> Footnotes -> Abbreviations LAST -> Diff highlight)
# ======================================================================================
def _one_tool_pipeline(file_path: str):
    if not file_path:
        return None

    name = os.path.basename(file_path)
    base, ext = os.path.splitext(name)
    ts = int(time.time())
    out_dir = "/mnt/data" if os.path.isdir("/mnt/data") else "."

    # Temp stage paths
    p_in    = os.path.join(out_dir, f"{base}_in_{ts}.docx")
    p_style = os.path.join(out_dir, f"{base}_style_{ts}.docx")
    p_foot  = os.path.join(out_dir, f"{base}_foot_{ts}.docx")
    p_abbr  = os.path.join(out_dir, f"{base}_abbr_{ts}.docx")
    out_path= os.path.join(out_dir, f"{base}_final_{ts}.docx")

    # Keep a frozen copy of the ORIGINAL (for diff)
    if ext.lower() == ".docx":
        p_orig = os.path.join(out_dir, f"{base}_orig_{ts}.docx")
        try:
            copyfile(file_path, p_orig)
        except Exception:
            p_orig = file_path  # fallback to reading original in-place
        cur = file_path
    else:
        text = _read_text_from_path(file_path)
        _save_text_as_docx(text, p_in)
        p_orig = p_in
        cur = p_in

    # Load modules
    test8_module, fprompt_module, ftool_module = _find_user_modules()

    # 1) STYLE
    try:
        process_docx_style(cur, p_style, model=DEFAULT_MODEL, max_tokens=MAX_TOKENS)
        cur = p_style
    except Exception:
        pass

    # 2) FOOTNOTES
    cur = _run_footnotes_step(ftool_module, fprompt_module, cur, p_foot)

    # 3) ABBREVIATIONS (test8 LAST)
    cur = _run_abbrev_step_last(test8_module, cur, p_abbr)

    # Finalise
    try:
        if cur != out_path:
            copyfile(cur, out_path)
        # 4) HIGHLIGHT CHANGES vs original (body + notes)
        try:
            highlight_doc_differences(p_orig, out_path)
        except Exception:
            pass
        return out_path
    except Exception:
        return cur if os.path.exists(cur) else None

# ======================================================================================
# GRADIO (single tool: upload -> download)
# ======================================================================================
with gr.Blocks(title="ISAS One-Tool Processor") as demo:
    gr.Markdown("### Upload a file (.docx or text) to get a fully processed .docx (changes highlighted in yellow)")
    inp = gr.File(label="Upload", file_count="single", type="filepath")
    out_file = gr.File(label="Download processed .docx")
    inp.change(_one_tool_pipeline, inputs=inp, outputs=out_file)

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=7860, share=True)
