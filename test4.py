# import os
# import anthropic
# import gradio as gr
# from docx import Document
# from dotenv import load_dotenv

# # Load environment variables from .env file
# load_dotenv("env (1)")  # Replace with actual path if different

# # Initialize Anthropic client using API key from env
# client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# # Define topic categories
# CATEGORIES = [
#     "Sustainability and the Environment",
#     "Strategic Technologies",
#     "Trade and Economics",
#     "Politics, Society and Governance",
#     "International Relations, Multipolarity and Multilateralism"
# ]

# # Extract title and first paragraph from a DOCX file
# def extract_intro_from_docx(file_path):
#     doc = Document(file_path)
#     paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
#     title = paragraphs[0] if paragraphs else ""
#     first_para = paragraphs[1] if len(paragraphs) > 1 else ""
#     return title, first_para

# # Classify the topic based on title and first paragraph
# def classify_topic_from_file(file_path):
#     title, paragraph = extract_intro_from_docx(file_path)
#     system_prompt = (
#         "You are an expert in classifying documents. Based only on the given title and paragraph, "
#         "classify the content into one of the following categories:\n"
#         + ", ".join(CATEGORIES) + ".\n"
#         "Respond ONLY with the exact name of the category. Do not explain or add anything else."
#     )
#     user_prompt = f"Title: {title}\nParagraph: {paragraph}"

#     response = client.messages.create(
#         model="claude-3-5-haiku-20241022",
#         max_tokens=100,
#         temperature=0.0,
#         system=system_prompt,
#         messages=[{"role": "user", "content": user_prompt}]
#     )

#     return response.content[0].text.strip()

# # Launch Gradio app
# demo = gr.Interface(
#     fn=classify_topic_from_file,
#     inputs=gr.File(label="Upload DOCX File", type="filepath"),
#     outputs=gr.Textbox(label="Predicted Topic"),
#     title="Topic Classifier with Claude",
#     description="Upload a .docx file. The app sends the title and first paragraph to Claude to determine the topic category."
# )

# demo.launch()
import os
import json
import re
import gradio as gr
import anthropic
from collections import defaultdict
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv("env (1)")

# Anthropic client
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

CATEGORY_TO_JSON = {
    "Sustainability and the Environment": "Sustainabilityandenvironment.json",
    "Strategic Technologies": "StatergicTechnologies.json",
    "Trade and Economics": "Trade.json",
    "Politics, Society and Governance": "PSG.json",
    "International Relations, Multipolarity and Multilateralism": "IEMTMLT.json"
}

def normalize_key(text):
    return re.sub(r'\s+', '', text.lower())

def clean_nested_expansion(text):
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\s*\((\w+)\)\)', r'\1 (\2)', text)
    text = re.sub(r'\b(\w[\w\s]+)\s*\(\1\)', r'\1', text)
    return text

def handle_special_case_usa(text):
    pattern1 = r'\b(The\s+United\s+States)\s*\(\s*US\s*\)\s+(of\s+America)\s*\(\s*USA\s*\)'
    text = re.sub(pattern1, r'\1 \2 (USA)', text, flags=re.IGNORECASE)
    pattern2 = r'\bThe United States\s+of America\b(?!\s*\(\s*USA\s*\))'
    text = re.sub(pattern2, r'The United States of America (USA)', text, flags=re.IGNORECASE)
    return text

def resolve_synonyms_first_mention_global(text, synonym_groups):
    for group in synonym_groups:
        positions = {}
        for term in group:
            match = re.search(rf'\b{re.escape(term)}\b', text, flags=re.IGNORECASE)
            if match:
                positions[term] = match.start()
        if positions:
            chosen = min(positions, key=positions.get)
            for alt in group:
                if alt.lower() != chosen.lower():
                    pattern = rf'\b{re.escape(alt)}\b'
                    text = re.sub(pattern, chosen, text, flags=re.IGNORECASE)
    return text

def classify_topic(doc_path):
    doc = Document(doc_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else ""
    paragraph = paragraphs[1] if len(paragraphs) > 1 else ""

    categories = list(CATEGORY_TO_JSON.keys())
    system_prompt = (
        "You are an expert in classifying documents. Based only on the given title and paragraph, "
        f"classify the content into one of the following categories:\n{', '.join(categories)}.\n"
        "Respond ONLY with the exact name of the category."
    )
    user_prompt = f"Title: {title}\nParagraph: {paragraph}"

    response = client.messages.create(
        model="claude-3-5-haiku-20241022",
        max_tokens=100,
        temperature=0,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}]
    )

    return response.content[0].text.strip()

def expand_abbreviations(doc_path, base_json_path="abbreviations.json"):
    with open(base_json_path, "r", encoding="utf-8") as f:
        base_map = json.load(f)

    category = classify_topic(doc_path)
    category_map = {}
    category_json = CATEGORY_TO_JSON.get(category)
    if category_json and os.path.exists(category_json):
        with open(category_json, "r", encoding="utf-8") as f:
            category_map = json.load(f)

    abbr_map = {**base_map, **category_map}
    counts = defaultdict(int)
    norm_map = {normalize_key(v): (k, v) for k, v in abbr_map.items()}

    input_doc = Document(doc_path)
    result_doc = Document()

    synonym_groups = [
        ["South Korea", "Republic of Korea", "ROK"],
        ["The United States", "The United States of America"],
        ["North Korea", "Democratic People's Republic of Korea", "DPRK"],
        ["China", "People Republic of China", "PRC"]
    ]

    def is_heading(paragraph):
        return paragraph.style.name.startswith("Heading")

    def is_in_table(paragraph):
        return paragraph._element.getparent().tag.endswith('tbl')

    header_texts = []
    footer_texts = []
    for section in input_doc.sections:
        header_texts.extend(p.text.strip() for p in section.header.paragraphs)
        footer_texts.extend(p.text.strip() for p in section.footer.paragraphs)

    def is_in_header_or_footer(text):
        return text.strip() in header_texts or text.strip() in footer_texts

    full_text = "\n".join([
        p.text for p in input_doc.paragraphs
        if not is_heading(p) and not is_in_table(p) and not is_in_header_or_footer(p.text)
    ])
    full_text = resolve_synonyms_first_mention_global(full_text, synonym_groups)

    non_skipped_texts = full_text.split("\n")
    idx = 0
    for para in input_doc.paragraphs:
        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(para.text):
            para.text = para.text
        else:
            para.text = non_skipped_texts[idx]
            idx += 1

    for para in input_doc.paragraphs:
        line = para.text

        if is_heading(para) or is_in_table(para) or is_in_header_or_footer(line):
            result_doc.add_paragraph(line)
            continue

        line = clean_nested_expansion(line)

        for norm_exp, (abbr, expansion) in norm_map.items():
            abbr_pattern = rf'\b{re.escape(abbr)}\b'
            exp_pattern = rf'\b{re.escape(expansion)}\b'
            patterns = [
                rf'{abbr}\s*\(\s*{expansion}\s*\)',
                rf'{expansion}\s*\(\s*{abbr}\s*\)',
                rf'{abbr}\s*\(\s*\w[\w\s]*\)',
                rf'\w[\w\s]*\(\s*{abbr}\s*\)',
                abbr_pattern,
                exp_pattern
            ]
            if any(re.search(p, line, re.IGNORECASE) for p in patterns):
                counts[abbr] += 1
                if counts[abbr] == 1:
                    correct_format = rf'{re.escape(expansion)}\s*\(\s*{re.escape(abbr)}\s*\)'
                    if not re.search(correct_format, line, re.IGNORECASE):
                        line = re.sub('|'.join(patterns), f"{expansion} ({abbr})", line, count=1, flags=re.IGNORECASE)
                else:
                    line = re.sub('|'.join(patterns), abbr, line, flags=re.IGNORECASE)

        line = handle_special_case_usa(line)
        result_doc.add_paragraph(line)

    out_path = "expanded_result.docx"
    result_doc.save(out_path)
    return out_path

def gradio_abbreviation_expand(doc_file):
    output_path = expand_abbreviations(doc_file.name)
    return "✅ Document processed with topic-specific and global abbreviations.", output_path

gr.Interface(
    fn=gradio_abbreviation_expand,
    inputs=[gr.File(label="Upload DOCX Document", file_types=[".docx"])],
    outputs=[gr.Text(label="Status"), gr.File(label="Download Processed File")],
    title="Abbreviation Expander with Topic-Specific Enhancement",
    description="Expands abbreviations using both a global base list and a category-specific list (auto-detected using Claude)."
).launch()
